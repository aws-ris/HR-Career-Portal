import os
from fastapi import FastAPI, Depends, HTTPException, status, BackgroundTasks, Response

from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, extract, or_
from database.database import engine, Base, get_db, SessionLocal

from database import models
import schemas
import datetime
from pydantic import BaseModel
from typing import List, Optional
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from utils.auth import generate_token, verify_token

MAX_RESUME_FILE_SIZE_BYTES = 5 * 1024 * 1024
ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_RESUME_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def validate_resume_upload(filename: str, content_type: Optional[str], size_bytes: int):
    name = (filename or "").lower()
    file_ext = os.path.splitext(name)[1]
    mime_type = (content_type or "").split(";")[0].strip().lower()

    if file_ext not in ALLOWED_RESUME_EXTENSIONS and mime_type not in ALLOWED_RESUME_MIME_TYPES:
        raise ValueError("Resume upload must be a PDF or DOCX file.")

    if size_bytes > MAX_RESUME_FILE_SIZE_BYTES:
        raise ValueError("Resume file must be 5MB or smaller.")

    return True


def get_resume_media_type(filename: str, fallback_content_type: Optional[str] = None) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "application/pdf"
    if name.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return (fallback_content_type or "application/octet-stream").split(";")[0].strip().lower()

auth_scheme = HTTPBearer(auto_error=False)

class LoginRequest(BaseModel):
    username: str
    password: str

def get_current_admin(
    token: Optional[str] = None,
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(auth_scheme)
):
    actual_token = None
    if credentials:
        actual_token = credentials.credentials
    elif token:
        actual_token = token
        
    if not actual_token:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing authentication credentials"
        )
        
    username = verify_token(actual_token)
    if not username:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired session token"
        )
    return username

# Note: In production, use Alembic migrations instead of create_all
# create_all is safe here — it only creates tables that don't exist yet
Base.metadata.create_all(bind=engine)

# Auto-migration: ensure key_terms column exists in job_postings table
try:
    from sqlalchemy import text
    with engine.connect() as conn:
        conn.execute(text("ALTER TABLE job_postings ADD COLUMN key_terms TEXT;"))
        conn.commit()
except Exception:
    pass

# Ensure default admin user exists in PostgreSQL DB
try:
    from utils.auth import hash_password
    with SessionLocal() as _db:
        _admin_user_env = os.getenv("HR_ADMIN_USERNAME", "hr_ris")
        _admin_pass_env = os.getenv("HR_ADMIN_PASSWORD", "ris@1234")
        _admin_rec = _db.query(models.AdminUser).filter(models.AdminUser.username == _admin_user_env).first()
        if not _admin_rec:
            _admin_rec = models.AdminUser(
                username=_admin_user_env,
                password_hash=hash_password(_admin_pass_env),
                full_name="HR Administrator"
            )
            _db.add(_admin_rec)
            _db.commit()
            print(f"[Admin Seed] Seeded default admin user '{_admin_user_env}' in database")
except Exception as _e:
    print(f"[Admin Seed Warning] Warning initializing admin user: {_e}")

app = FastAPI(title="RIS Hiring Portal API", version="2.0.0")

import logging
from fastapi.responses import JSONResponse

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ris_portal")

@app.exception_handler(Exception)
async def global_exception_handler(request, exc: Exception):
    logger.error(f"Unhandled server exception on {request.method} {request.url}: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "status": "error",
            "message": "An internal processing issue occurred. For assistance, email parmod.kumar@ris.org.in.",
            "detail": "Internal Processing Error"
        }
    )

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

@app.post("/api/v1/auth/login")
@app.post("/v1/auth/login")
@app.post("/auth/login")
def login_admin(req: LoginRequest, db: Session = Depends(get_db)):
    from utils.auth import verify_password
    admin_user_env = os.getenv("HR_ADMIN_USERNAME", "hr_ris")
    admin_pass_env = os.getenv("HR_ADMIN_PASSWORD", "ris@1234")
    
    # 1. Database-backed authentication (Sole authority if user exists in DB)
    db_admin = db.query(models.AdminUser).filter(
        models.AdminUser.username == req.username,
        models.AdminUser.is_active == True
    ).first()
    
    if db_admin:
        if verify_password(req.password, db_admin.password_hash):
            db_admin.last_login_at = datetime.datetime.utcnow()
            db.commit()
            token = generate_token(db_admin.username)
            return {
                "status": "success",
                "token": token,
                "username": db_admin.username,
                "full_name": db_admin.full_name or "HR Administrator"
            }
        else:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid username or password"
            )
        
    # 2. Environment variable fallback ONLY if user record does not exist in DB yet
    if req.username == admin_user_env and req.password == admin_pass_env:
        token = generate_token(req.username)
        return {"status": "success", "token": token, "username": admin_user_env, "full_name": "HR Administrator"}
        
    raise HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Invalid username or password"
    )

@app.post("/api/v1/auth/change-password")
def change_admin_password(
    req: ChangePasswordRequest,
    current_username: str = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    from utils.auth import verify_password, hash_password
    if not req.new_password or len(req.new_password.strip()) < 6:
        raise HTTPException(status_code=400, detail="New password must be at least 6 characters long.")

    db_admin = db.query(models.AdminUser).filter(models.AdminUser.username == current_username).first()
    admin_pass_env = os.getenv("HR_ADMIN_PASSWORD", "ris@1234")

    if db_admin:
        if not verify_password(req.old_password, db_admin.password_hash):
            raise HTTPException(status_code=400, detail="Current password is incorrect.")
        db_admin.password_hash = hash_password(req.new_password)
        db_admin.updated_at = datetime.datetime.utcnow()
    else:
        if req.old_password != admin_pass_env:
            raise HTTPException(status_code=400, detail="Current password is incorrect.")
        db_admin = models.AdminUser(
            username=current_username,
            password_hash=hash_password(req.new_password),
            full_name="HR Administrator"
        )
        db.add(db_admin)

    db.commit()
    return {"status": "success", "message": "Password updated successfully!"}

from fastapi import File, UploadFile
from fastapi.responses import FileResponse
import os

# AI Preloading disabled - using cloud API



@app.on_event("startup")
def startup_migration():
    """
    Automatically check and apply database schema migrations at startup.
    This runs every time the server boots, ensuring zero manual intervention.
    """
    from database.database import SessionLocal
    from sqlalchemy import text
    from database import models
    import datetime
    
    db = SessionLocal()
    try:
        print("Running automatic database schema migrations...")
        
        # 1. Add missing columns safely for both SQLite and PostgreSQL
        is_sqlite = "sqlite" in str(db.bind.url)
        
        column_migrations = [
            ("job_postings", "min_pay", "INTEGER"),
            ("job_postings", "max_pay", "INTEGER"),
            ("job_postings", "min_experience", "INTEGER"),
            ("job_postings", "max_experience", "INTEGER"),
            ("job_postings", "contract_period", "INTEGER"),
            ("job_postings", "job_mode", "VARCHAR(100)"),
            ("job_postings", "pay_band", "VARCHAR(50)"),
            ("job_postings", "pay_level", "VARCHAR(50)"),
            
            ("candidate_metadata", "country_code", "VARCHAR(10)"),
            ("candidate_metadata", "nationality", "VARCHAR(100) DEFAULT 'Indian'"),
            ("candidate_metadata", "pincode", "VARCHAR(20)"),
            ("candidate_metadata", "is_international_address", "BOOLEAN DEFAULT FALSE"),
            ("candidate_metadata", "international_address", "TEXT"),
            ("candidate_metadata", "age", "INTEGER"),
            ("candidate_metadata", "city", "VARCHAR(100)"),
            ("candidate_metadata", "last_salary", "FLOAT"),

            ("candidate_higher_education", "phd_domain", "VARCHAR(255)"),
            ("candidate_metadata", "worked_at_ris", "BOOLEAN DEFAULT FALSE"),
            ("candidate_metadata", "ris_designation", "VARCHAR(200)"),
            ("candidate_metadata", "ris_start_date", "DATE"),
            ("candidate_metadata", "ris_end_date", "DATE"),
            ("candidate_metadata", "ris_is_current", "BOOLEAN DEFAULT FALSE"),

            
            ("candidate_links_about", "about", "TEXT"),
            ("candidate_links_about", "sop", "TEXT"),
            ("candidate_links_about", "extracurriculars", "TEXT"),
            ("candidate_links_about", "pub_books", "INTEGER DEFAULT 0"),
            ("candidate_links_about", "pub_papers", "INTEGER DEFAULT 0"),
            ("candidate_links_about", "pub_chapters", "INTEGER DEFAULT 0"),
            ("candidate_links_about", "pub_reports", "INTEGER DEFAULT 0"),
            ("candidate_links_about", "pub_policy_briefs", "INTEGER DEFAULT 0"),
            ("candidate_links_about", "how_heard", "VARCHAR(500)"),
            
            ("candidate_resume_payload", "pdf_blob", "BLOB" if is_sqlite else "BYTEA"),
            ("application_tracking", "profile_score", "FLOAT"),
            ("candidate_schooling", "class_x_year", "INTEGER"),
            ("candidate_schooling", "class_xii_year", "INTEGER"),
            
            ("candidate_higher_education", "is_pursuing", "BOOLEAN DEFAULT FALSE"),
            ("candidate_higher_education", "duration_value", "INTEGER"),
            ("candidate_higher_education", "duration_unit", "VARCHAR(10)")
        ]
        
        if is_sqlite:
            for table, col, col_type in column_migrations:
                try:
                    res = db.execute(text(f"PRAGMA table_info({table});")).fetchall()
                    existing_cols = [r[1] for r in res]
                    if col not in existing_cols:
                        db.execute(text(f"ALTER TABLE {table} ADD COLUMN {col} {col_type};"))
                        db.commit()
                        print(f"SQLite added column {col} to {table}")
                except Exception as sq_err:
                    db.rollback()
                    print(f"SQLite migration error for {table}.{col}: {sq_err}")
        else:
            for table, col, col_type in column_migrations:
                try:
                    db.execute(text(f"ALTER TABLE {table} ADD COLUMN IF NOT EXISTS {col} {col_type};"))
                    db.commit()
                except Exception as pg_err:
                    db.rollback()
                    print(f"Postgres migration error for {table}.{col}: {pg_err}")
            
            # Automatically drop the category column from postgres
            try:
                db.execute(text("ALTER TABLE candidate_metadata DROP COLUMN IF EXISTS category;"))
                db.commit()
                print("Postgres safely dropped category column.")
            except Exception as pg_err:
                db.rollback()
                print(f"Postgres drop error for category: {pg_err}")
        
        # 2. Automatically backfill candidate ages and categories
        try:
            candidates = db.query(models.CandidateMetadata).all()
            today = datetime.date.today()
            sample_cats = ["General (UR)", "OBC (Non-Creamy Layer)", "SC (Scheduled Caste)", "ST (Scheduled Tribe)", "EWS (Economically Weaker Section)"]
            backfilled_count = 0
            cat_count = 0
            for idx, cand in enumerate(candidates):
                if cand.dob and cand.age is None:
                    cand.age = today.year - cand.dob.year - ((today.month, today.day) < (cand.dob.month, cand.dob.day))
                    backfilled_count += 1

            if backfilled_count > 0 or cat_count > 0:
                db.commit()
                print(f"Auto-backfilled ages for {backfilled_count} and categories for {cat_count} candidates.")
        except Exception as e:
            print(f"Error backfilling candidate metadata: {e}")
            db.rollback()

        # 3. Auto-seed if the compliant G20 Professor job doesn't exist
        try:
            target_job = db.query(models.JobPosting).filter(
                models.JobPosting.title == "Professor (Trade & Investment Policy)"
            ).first()
            if not target_job:
                print("New G20/Blue Economy jobs not found in database. Triggering automatic database re-seed...")
                from scratch.seed_final_portal_data import main as run_seeder
                run_seeder()
        except Exception as e:
            print(f"Error checking/running automatic database seeder: {e}")
            
        print("Automatic database migrations completed successfully.")
    except Exception as e:
        print(f"Automatic migration critical error: {e}")
    finally:
        db.close()


# ─────────────────────────────────────────────
# Public Job Board Access
# ─────────────────────────────────────────────

# ─────────────────────────────────────────────
# Public Job Board Access

# ─────────────────────────────────────────────
# Public Job Board Access

# ─────────────────────────────────────────────
# Public Job Board Access
# ─────────────────────────────────────────────

@app.get("/api/v1/public/jobs")
@app.get("/v1/public/jobs")
@app.get("/public/jobs")
def get_public_jobs(db: Session = Depends(get_db)):
    """
    Returns only 'open' jobs whose deadline has not passed.
    Automatically marks expired jobs as 'closed'.
    """
    today = datetime.date.today()

    # Auto-close any open jobs whose deadline has passed
    expired_jobs = db.query(models.JobPosting).filter(
        models.JobPosting.status == 'open',
        models.JobPosting.deadline != None,
        models.JobPosting.deadline < today,
        models.JobPosting.is_deleted == False
    ).all()
    if expired_jobs:
        for ej in expired_jobs:
            ej.status = 'closed'
        db.commit()

    jobs = db.query(models.JobPosting).filter(
        models.JobPosting.status == 'open',
        models.JobPosting.is_deleted == False,
        or_(models.JobPosting.deadline == None, models.JobPosting.deadline >= today)
    ).order_by(models.JobPosting.deadline.asc()).all()
    
    return [{
        "id": j.id,
        "title": j.title,
        "position": j.position,
        "division": j.division,
        "location": j.location,
        "deadline": j.deadline,
        "description": j.description,
        "key_terms": j.key_terms,
        "requirements": j.requirements,
        "min_pay": j.min_pay,
        "max_pay": j.max_pay,
        "min_experience": j.min_experience,
        "max_experience": j.max_experience,
        "contract_period": j.contract_period,
        "job_mode": j.job_mode,
        "created_at": j.created_at.isoformat() if j.created_at else None
    } for j in jobs]

@app.get("/api/v1/public/jobs/{job_id}")
def get_public_job_detail(job_id: str, db: Session = Depends(get_db)):
    """
    Returns full public details for a specific job if it is open and deadline is valid.
    """
    today = datetime.date.today()
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.is_deleted == False
    ).first()
    
    if not job:
        raise HTTPException(status_code=404, detail="Job posting not found")

    if job.status == 'open' and job.deadline and job.deadline < today:
        job.status = 'closed'
        db.commit()

    if job.status != 'open' or (job.deadline and job.deadline < today):
        raise HTTPException(status_code=400, detail="This job vacancy is closed or no longer accepting applications.")
        
    return {
        "id": job.id,
        "title": job.title,
        "position": job.position,
        "division": job.division,
        "description": job.description,
        "key_terms": job.key_terms,
        "requirements": job.requirements,
        "deadline": job.deadline,
        "min_pay": job.min_pay,
        "max_pay": job.max_pay,
        "min_experience": job.min_experience,
        "max_experience": job.max_experience,
        "contract_period": job.contract_period,
        "job_mode": job.job_mode,
        "created_at": job.created_at.isoformat() if job.created_at else None
    }


def async_score_candidate_bg(candidate_id: str, app_tracking_id: str, job_id: Optional[str]):
    db = SessionLocal()
    try:
        candidate = db.query(models.CandidateMetadata).filter(models.CandidateMetadata.id == candidate_id).first()
        app_tracking = db.query(models.ApplicationTracking).filter(models.ApplicationTracking.id == app_tracking_id).first()
        if candidate and app_tracking:
            job_posting = db.query(models.JobPosting).filter(models.JobPosting.id == job_id).first() if job_id else None
            min_exp = job_posting.min_experience if job_posting else 1.0
            from utils.scoring import calculate_candidate_score
            score_res = calculate_candidate_score(candidate, min_exp)
            app_tracking.profile_score = score_res["total_score"]
            db.commit()
    except Exception as e:
        db.rollback()
        print(f"[Async Scoring Error] {e}")
    finally:
        db.close()

# ─────────────────────────────────────────────
# Candidate Application Submission
# ─────────────────────────────────────────────
@app.post(
    "/api/v1/applications",
    response_model=schemas.CandidateResponse,
    status_code=status.HTTP_201_CREATED
)
def create_application(payload: schemas.CandidateCreate, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    # Validate that the job is open and deadline has not passed
    if payload.job_id:
        today = datetime.date.today()
        job = db.query(models.JobPosting).filter(
            models.JobPosting.id == payload.job_id,
            models.JobPosting.is_deleted == False
        ).first()

        if job and job.status == 'open' and job.deadline and job.deadline < today:
            job.status = 'closed'
            db.commit()

        if not job or job.status != 'open' or (job.deadline and job.deadline < today):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="This job vacancy is closed and is no longer accepting applications."
            )


    try:
        # Check if candidate email already exists (Error prevention & hassle-free submission)
        candidate = db.query(models.CandidateMetadata).filter(
            models.CandidateMetadata.email == payload.email
        ).first()
        
        is_new_candidate = False
        if candidate:
            # Update existing candidate details
            candidate.full_name = payload.full_name
            candidate.nationality = getattr(payload, 'nationality', None) or 'Indian'
            candidate.country_code = payload.country_code
            candidate.mobile_no = payload.mobile_no
            candidate.dob = payload.dob
            candidate.gender = payload.gender
            candidate.city = payload.city
            candidate.state = payload.state
            candidate.pincode = payload.pincode
            candidate.is_international_address = getattr(payload, 'is_international_address', False) or False
            candidate.international_address = getattr(payload, 'international_address', None)
            candidate.years_of_experience = payload.years_of_experience
            candidate.last_salary = payload.last_salary
            candidate.worked_at_ris = getattr(payload, 'worked_at_ris', False) or False
            candidate.ris_designation = getattr(payload, 'ris_designation', None)
            candidate.ris_start_date = getattr(payload, 'ris_start_date', None)
            candidate.ris_end_date = getattr(payload, 'ris_end_date', None)
            candidate.ris_is_current = getattr(payload, 'ris_is_current', False) or False
            
            # Clean up old relations to prevent duplicates
            if candidate.schooling:
                db.delete(candidate.schooling)
            if candidate.links_about:
                db.delete(candidate.links_about)
            
            db.query(models.CandidateHigherEducation).filter_by(candidate_id=candidate.id).delete()
            db.query(models.CandidatePublication).filter_by(candidate_id=candidate.id).delete()
            db.query(models.CandidateWorkExperience).filter_by(candidate_id=candidate.id).delete()
            db.flush()
        else:
            is_new_candidate = True
            # Create new candidate record
            candidate = models.CandidateMetadata(
                full_name           = payload.full_name,
                email               = payload.email,
                nationality         = getattr(payload, 'nationality', None) or 'Indian',
                country_code        = payload.country_code,
                mobile_no           = payload.mobile_no,
                dob                 = payload.dob,
                gender              = payload.gender,
                city                = payload.city,
                state               = payload.state,
                pincode             = payload.pincode,
                is_international_address = getattr(payload, 'is_international_address', False) or False,
                international_address = getattr(payload, 'international_address', None),
                years_of_experience = payload.years_of_experience,
                last_salary         = payload.last_salary,
                worked_at_ris       = getattr(payload, 'worked_at_ris', False) or False,
                ris_designation     = getattr(payload, 'ris_designation', None),
                ris_start_date      = getattr(payload, 'ris_start_date', None),
                ris_end_date        = getattr(payload, 'ris_end_date', None),
                ris_is_current      = getattr(payload, 'ris_is_current', False) or False
            )

            db.add(candidate)
            db.flush()

        # Create or update application tracking record
        app_tracking = None
        if not is_new_candidate:
            app_tracking = db.query(models.ApplicationTracking).filter_by(
                candidate_id=candidate.id,
                job_id=payload.job_id
            ).first()
            
        if app_tracking:
            app_tracking.position_applied = payload.position_applied
            app_tracking.admin_department = payload.admin_department
            app_tracking.current_status = 'received'
            app_tracking.updated_at = datetime.datetime.utcnow()
        else:
            app_tracking = models.ApplicationTracking(
                candidate_id     = candidate.id,
                job_id           = payload.job_id,
                position_applied = payload.position_applied,
                admin_department = payload.admin_department,
                current_status   = 'received'
            )
            db.add(app_tracking)
        db.flush()

        # Re-add relations (schooling, links_about, higher_education, publications, work_experiences)
        # 3. Candidate Links & About
        links_about = models.CandidateLinksAbout(
            candidate_id   = candidate.id,
            about          = payload.about,
            sop            = payload.sop,
            google_scholar = payload.google_scholar,
            linkedin       = payload.linkedin,
            pub_books      = payload.pub_books,
            pub_papers     = payload.pub_papers,
            pub_chapters   = payload.pub_chapters,
            pub_reports    = payload.pub_reports,
            pub_policy_briefs = payload.pub_policy_briefs,
            how_heard      = getattr(payload, 'how_heard', None)
        )
        db.add(links_about)

        # 4. Schooling (1:1)
        db.add(models.CandidateSchooling(
            candidate_id          = candidate.id,
            class_x_school        = payload.schooling.class_x_school,
            class_x_board         = payload.schooling.class_x_board,
            class_x_score_type    = payload.schooling.class_x_score_type.value if hasattr(payload.schooling.class_x_score_type, 'value') else payload.schooling.class_x_score_type,
            class_x_score_value   = payload.schooling.class_x_score_value,
            class_x_year          = payload.schooling.class_x_year,
            class_xii_school       = payload.schooling.class_xii_school,
            class_xii_board        = payload.schooling.class_xii_board,
            class_xii_score_type   = payload.schooling.class_xii_score_type.value if hasattr(payload.schooling.class_xii_score_type, 'value') else payload.schooling.class_xii_score_type,
            class_xii_score_value  = payload.schooling.class_xii_score_value,
            class_xii_year         = payload.schooling.class_xii_year,
        ))

        # 5. Higher Education (1:N)
        for edu in payload.higher_education:
            db.add(models.CandidateHigherEducation(
                candidate_id   = candidate.id,
                level          = edu.level,
                university     = edu.university,
                degree_name    = edu.degree_name,
                score_type     = edu.score_type,
                score_value    = edu.score_value,
                grad_year      = edu.grad_year,
                is_pursuing    = getattr(edu, 'is_pursuing', False) or False,
                duration_value = getattr(edu, 'duration_value', None),
                duration_unit  = getattr(edu, 'duration_unit', None),
                entry_order    = edu.entry_order,
            ))

        # 6. Publications (1:N)
        for pub in payload.publications:
            db.add(models.CandidatePublication(
                candidate_id = candidate.id,
                pub_type     = pub.pub_type,
                title        = pub.title,
                parent_book  = pub.parent_book,
                entry_order  = pub.entry_order,
            ))

        # 7. Work Experiences (1:N)
        for w in payload.work_experiences:
            db.add(models.CandidateWorkExperience(
                candidate_id = candidate.id,
                company_name = w.company_name,
                role         = w.role,
                start_date   = w.start_date,
                end_date     = w.end_date,
                is_current   = w.is_current,
                entry_order  = w.entry_order,
            ))

        db.flush()

        # 8. Seed status history
        db.add(models.ApplicationStatusHistory(
            application_tracking_id = app_tracking.id,
            status                 = 'received',
            changed_by             = 'SYSTEM',
            notes                  = 'Application updated/resubmitted' if not is_new_candidate else 'Application submitted',
        ))

        db.commit()
        db.refresh(candidate)
        
        # 9. Offload AI Scoring & Tokenization to BackgroundTasks (Ultra-fast <30ms submission!)
        background_tasks.add_task(async_score_candidate_bg, candidate.id, app_tracking.id, payload.job_id)
        background_tasks.add_task(tokenize_candidate_bg, candidate.id, payload.job_id)
        
        # 10. Trigger n8n webhook event in background (asynchronous non-blocking)
        from utils.webhooks import trigger_n8n_event
        submitted_at_formatted = datetime.datetime.now().strftime("%d-%m-%Y")
        trigger_n8n_event(background_tasks, "candidate-applied", {
            "candidate_id": candidate.id,
            "application_id": app_tracking.id if app_tracking else candidate.id,
            "full_name": candidate.full_name,
            "email": candidate.email,
            "mobile_no": candidate.mobile_no,
            "position_applied": payload.position_applied,
            "admin_department": payload.admin_department,
            "job_id": payload.job_id,
            "submitted_at": submitted_at_formatted,
            "how_heard": getattr(payload, 'how_heard', None)
        })

        return candidate
    except Exception as e:
        db.rollback()
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=str(e))


def tokenize_candidate(db: Session, candidate: models.CandidateMetadata):
    """
    Tokenizes candidate education, work experience, and publications to populate the TokenRegistry.
    This provides autocomplete suggestions for the HR filter interface.
    """
    # Get all applications to know which jobs to register these tokens for
    applications = db.query(models.ApplicationTracking).filter(
        models.ApplicationTracking.candidate_id == candidate.id
    ).all()
    
    job_ids = [app.job_id for app in applications if app.job_id]
    if not job_ids:
        return

    # Extract raw values and map to token types
    tokens_to_add = []
    
    # 1. Higher Education
    if candidate.higher_education:
        for edu in candidate.higher_education:
            if edu.university:
                tokens_to_add.append(('university', edu.university))
            if edu.degree_name:
                tokens_to_add.append(('degree', edu.degree_name))
                
    # 2. Work Experience
    if candidate.work_experiences:
        for exp in candidate.work_experiences:
            if exp.company_name:
                tokens_to_add.append(('company', exp.company_name))
            if exp.role:
                tokens_to_add.append(('role', exp.role))
                
    # 3. Publications
    if candidate.publications:
        for pub in candidate.publications:
            if pub.title:
                tokens_to_add.append(('pub_title', pub.title))

    for job_id in job_ids:
        for token_type, value in tokens_to_add:
            if not value or not value.strip():
                continue
            val_clean = value.strip()
            val_norm = val_clean.lower()
            
            # Check if this token already exists for this job and type
            existing = db.query(models.TokenRegistry).filter(
                models.TokenRegistry.job_id == job_id,
                models.TokenRegistry.token_type == token_type,
                models.TokenRegistry.normalized == val_norm
            ).first()
            
            if existing:
                existing.frequency += 1
            else:
                db.add(models.TokenRegistry(
                    job_id=job_id,
                    token_type=token_type,
                    token_value=val_clean,
                    normalized=val_norm,
                    frequency=1
                ))
    try:
        db.commit()
    except Exception as e:
        db.rollback()
        print(f"Error saving tokens in tokenize_candidate: {e}")


def tokenize_candidate_bg(candidate_id: str, job_id: Optional[str] = None):
    """
    Background wrapper to tokenize candidate structured data.
    """
    from database.database import SessionLocal
    db = SessionLocal()
    try:
        candidate = db.query(models.CandidateMetadata).filter(models.CandidateMetadata.id == candidate_id).first()
        if candidate:
            tokenize_candidate(db, candidate)
    except Exception as e:
        print(f"Background tokenization error for {candidate_id}: {e}")
    finally:
        db.close()


# ─────────────────────────────────────────────
# Get Full Candidate Profile
# ─────────────────────────────────────────────
@app.get(
    "/api/v1/applications/{candidate_id}",
    response_model=schemas.CandidateFullResponse
)
def get_application(candidate_id: str, db: Session = Depends(get_db)):
    candidate = db.query(models.CandidateMetadata).filter(
        models.CandidateMetadata.id == candidate_id
    ).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Application not found")
    return candidate


# ─────────────────────────────────────────────
# Application Status Tracker (public)
# ─────────────────────────────────────────────
@app.get("/api/v1/applications/{candidate_id}/status")
def get_status(candidate_id: str, db: Session = Depends(get_db)):
    return {
        "id":               candidate.id,
        "full_name":        candidate.full_name,
        "position_applied": candidate.position_applied,
        "current_status":   candidate.current_status,
        "submitted_at":     candidate.submitted_at,
    }


# ─────────────────────────────────────────────
# Job Postings (HR)
# ─────────────────────────────────────────────


@app.post(
    "/api/v1/jobs",
    dependencies=[Depends(get_current_admin)],
    response_model=schemas.JobPostingResponse,
    status_code=status.HTTP_201_CREATED
)
def create_job(payload: schemas.JobPostingCreate, db: Session = Depends(get_db)):
    job = models.JobPosting(**payload.model_dump())
    db.add(job)
    try:
        db.commit()
        db.refresh(job)
        return job
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/v1/jobs/{job_id}", response_model=schemas.JobPostingResponse, dependencies=[Depends(get_current_admin)])
def get_job(job_id: str, db: Session = Depends(get_db)):
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.is_deleted == False
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job posting not found")
    return job


# ─────────────────────────────────────────────
# HR Stats (KPI Cards)
# ─────────────────────────────────────────────
@app.get("/api/v1/hr/stats", dependencies=[Depends(get_current_admin)])
def get_hr_stats(db: Session = Depends(get_db)):
    current_year = datetime.datetime.now().year
    today = datetime.date.today()
    week_later = today + datetime.timedelta(days=7)

    open_positions = db.query(models.JobPosting).filter(
        models.JobPosting.status == 'open',
        models.JobPosting.is_deleted == False
    ).count()

    total_applicants_year = db.query(models.ApplicationTracking).filter(
        extract('year', models.ApplicationTracking.submitted_at) == current_year
    ).count()

    closing_soon = db.query(models.JobPosting).filter(
        models.JobPosting.status == 'open',
        models.JobPosting.is_deleted == False,
        models.JobPosting.deadline != None,
        models.JobPosting.deadline >= today,
        models.JobPosting.deadline <= week_later
    ).count()

    return {
        "open_positions": open_positions,
        "total_applicants_year": total_applicants_year,
        "closing_soon": closing_soon,
    }

@app.get("/api/v1/hr/analytics/global", dependencies=[Depends(get_current_admin)])
def get_global_analytics(db: Session = Depends(get_db)):
    gender_stats = db.query(
        models.CandidateMetadata.gender, 
        func.count(models.CandidateMetadata.id)
    ).group_by(models.CandidateMetadata.gender).all()

    state_stats = db.query(
        models.CandidateMetadata.state, 
        func.count(models.CandidateMetadata.id)
    ).group_by(models.CandidateMetadata.state).order_by(func.count(models.CandidateMetadata.id).desc()).limit(5).all()

    candidate_records = db.query(models.CandidateMetadata).options(
        joinedload(models.CandidateMetadata.higher_education)
    ).all()

    phd_count = 0
    pg_count = 0
    ug_count = 0

    PHD_LEVELS = {'phd', 'doctorate', 'ph.d', 'ph.d.', 'doctoral'}
    PG_LEVELS = {'postgrad', 'masters', 'master', 'post-graduate', 'post graduate', 'pg', 'postgraduate'}
    UG_LEVELS = {'undergrad', 'bachelors', 'bachelor', 'under-graduate', 'under graduate', 'ug', 'undergraduate'}

    for cand in candidate_records:
        levels = set(e.level.lower() for e in (cand.higher_education or []) if e.level)
        if levels & PHD_LEVELS:
            phd_count += 1
        elif levels & PG_LEVELS:
            pg_count += 1
        elif levels & UG_LEVELS or (cand.higher_education and len(cand.higher_education) > 0):
            ug_count += 1

    return {
        "gender": [{"name": g if g else "Other", "value": c} for g, c in gender_stats],
        "states": [{"name": s if s else "Unknown", "value": c} for s, c in state_stats],
        "education": [
            {"name": "PhD", "value": phd_count},
            {"name": "Masters", "value": pg_count},
            {"name": "Bachelors", "value": ug_count}
        ]
    }

@app.get("/api/v1/jobs/{job_id}/analytics", dependencies=[Depends(get_current_admin)])
def get_job_analytics(job_id: str, db: Session = Depends(get_db)):
    clean_id = str(job_id).strip()

    gender_stats = db.query(
        models.CandidateMetadata.gender, 
        func.count(models.CandidateMetadata.id)
    ).join(models.ApplicationTracking).filter(
        models.ApplicationTracking.job_id == clean_id
    ).group_by(models.CandidateMetadata.gender).all()

    state_stats = db.query(
        models.CandidateMetadata.state, 
        func.count(models.CandidateMetadata.id)
    ).join(models.ApplicationTracking).filter(
        models.ApplicationTracking.job_id == clean_id
    ).group_by(models.CandidateMetadata.state).order_by(func.count(models.CandidateMetadata.id).desc()).limit(5).all()

    candidate_records = db.query(models.CandidateMetadata).options(
        joinedload(models.CandidateMetadata.higher_education)
    ).join(models.ApplicationTracking).filter(
        models.ApplicationTracking.job_id == clean_id
    ).all()

    phd_count = 0
    pg_count = 0
    ug_count = 0

    PHD_LEVELS = {'phd', 'doctorate', 'ph.d', 'ph.d.', 'doctoral'}
    PG_LEVELS = {'postgrad', 'masters', 'master', 'post-graduate', 'post graduate', 'pg', 'postgraduate'}
    UG_LEVELS = {'undergrad', 'bachelors', 'bachelor', 'under-graduate', 'under graduate', 'ug', 'undergraduate'}

    for cand in candidate_records:
        levels = set(e.level.lower() for e in (cand.higher_education or []) if e.level)
        if levels & PHD_LEVELS:
            phd_count += 1
        elif levels & PG_LEVELS:
            pg_count += 1
        elif levels & UG_LEVELS or (cand.higher_education and len(cand.higher_education) > 0):
            ug_count += 1

    return {
        "gender": [{"name": g if g else "Other", "value": c} for g, c in gender_stats],
        "states": [{"name": s if s else "Unknown", "value": c} for s, c in state_stats],
        "education": [
            {"name": "PhD", "value": phd_count},
            {"name": "Masters", "value": pg_count},
            {"name": "Bachelors", "value": ug_count}
        ]
    }

@app.get("/api/v1/candidates/{candidate_id}/full_profile", dependencies=[Depends(get_current_admin)])
def get_full_profile(candidate_id: str, job_id: Optional[str] = None, db: Session = Depends(get_db)):
    candidate = db.query(models.CandidateMetadata).options(
        joinedload(models.CandidateMetadata.schooling),
        joinedload(models.CandidateMetadata.higher_education),
        joinedload(models.CandidateMetadata.publications),
        joinedload(models.CandidateMetadata.work_experiences),
        joinedload(models.CandidateMetadata.links_about),
        joinedload(models.CandidateMetadata.applications),
    ).filter(
        models.CandidateMetadata.id == candidate_id
    ).first()
    
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")

    min_exp = 1.0
    if job_id:
        job_posting = db.query(models.JobPosting).filter(models.JobPosting.id == job_id).first()
        if job_posting and job_posting.min_experience is not None:
            min_exp = job_posting.min_experience
    else:
        # Fallback: check if they have applied to any job, use the latest application's min_experience
        latest_app = db.query(models.ApplicationTracking).filter(
            models.ApplicationTracking.candidate_id == candidate_id
        ).order_by(models.ApplicationTracking.submitted_at.desc()).first()
        if latest_app and latest_app.job_id:
            job_posting = db.query(models.JobPosting).filter(models.JobPosting.id == latest_app.job_id).first()
            if job_posting and job_posting.min_experience is not None:
                min_exp = job_posting.min_experience

    from utils.scoring import calculate_candidate_score
    score_res = calculate_candidate_score(candidate, min_exp)

    schooling_data = None
    if candidate.schooling:
        schooling_data = {
            "class_x_school": candidate.schooling.class_x_school,
            "class_x_board": candidate.schooling.class_x_board,
            "class_x_score_type": candidate.schooling.class_x_score_type,
            "class_x_score_value": candidate.schooling.class_x_score_value,
            "class_x_year": getattr(candidate.schooling, "class_x_year", None),
            "class_xii_school": candidate.schooling.class_xii_school,
            "class_xii_board": candidate.schooling.class_xii_board,
            "class_xii_score_type": candidate.schooling.class_xii_score_type,
            "class_xii_score_value": candidate.schooling.class_xii_score_value,
            "class_xii_year": getattr(candidate.schooling, "class_xii_year", None)
        }

    # Format higher education by level
    grad = [e for e in candidate.higher_education if e.level == 'undergrad']
    postgrad = [e for e in candidate.higher_education if e.level == 'postgrad']
    phd = [e for e in candidate.higher_education if e.level == 'phd']

    res_profile = {
        "id": candidate.id,
        "full_name": candidate.full_name,
        "email": candidate.email,
        "nationality": getattr(candidate, "nationality", None) or "Indian",
        "country_code": candidate.country_code,
        "mobile_no": candidate.mobile_no,
        "dob": candidate.dob.isoformat() if candidate.dob else None,
        "gender": candidate.gender,
        "state": candidate.state,
        "city": candidate.city,
        "pincode": candidate.pincode,
        "is_international_address": getattr(candidate, "is_international_address", False) or False,
        "international_address": getattr(candidate, "international_address", None),

        "age": candidate.age or (
            (datetime.date.today().year - candidate.dob.year - ((datetime.date.today().month, datetime.date.today().day) < (candidate.dob.month, candidate.dob.day)))
            if candidate.dob else None
        ),
        "years_of_experience": candidate.years_of_experience,
        "worked_at_ris": getattr(candidate, "worked_at_ris", False) or False,
        "ris_designation": getattr(candidate, "ris_designation", None),
        "ris_start_date": candidate.ris_start_date.isoformat() if getattr(candidate, "ris_start_date", None) else None,
        "ris_end_date": candidate.ris_end_date.isoformat() if getattr(candidate, "ris_end_date", None) else None,
        "ris_is_current": getattr(candidate, "ris_is_current", False) or False,
        "profile_score": score_res["total_score"],
        "profile_score_breakdown": score_res["breakdown"],
        "about": candidate.links_about.about if candidate.links_about else None,
        "sop": candidate.links_about.sop if candidate.links_about else None,
        "google_scholar": candidate.links_about.google_scholar if candidate.links_about else None,
        "linkedin": candidate.links_about.linkedin if candidate.links_about else None,
        "pub_books": candidate.links_about.pub_books if candidate.links_about else 0,
        "pub_papers": candidate.links_about.pub_papers if candidate.links_about else 0,
        "pub_chapters": candidate.links_about.pub_chapters if candidate.links_about else 0,
        "pub_reports": candidate.links_about.pub_reports if candidate.links_about else 0,
        "pub_policy_briefs": candidate.links_about.pub_policy_briefs if candidate.links_about else 0,
        "schooling": schooling_data,
        "graduation": [{"degree_name": g.degree_name, "university": g.university, "score_type": g.score_type, "score_value": g.score_value, "grad_year": g.grad_year} for g in grad],
        "postgraduate": [{"degree_name": p.degree_name, "university": p.university, "score_type": p.score_type, "score_value": p.score_value, "grad_year": p.grad_year} for p in postgrad],
        "doctorate": [{"university": d.university, "thesis_title": d.degree_name, "phd_domain": getattr(d, "phd_domain", None), "grad_year": d.grad_year, "is_pursuing": getattr(d, "is_pursuing", False)} for d in phd],
        "work_experiences": [{"role": w.role, "company_name": w.company_name, "start_date": w.start_date.isoformat() if w.start_date else None, "end_date": w.end_date.isoformat() if w.end_date else None, "is_current": w.is_current} for w in candidate.work_experiences],
        "publications": [{"pub_type": pub.pub_type, "title": pub.title, "parent_book": pub.parent_book} for pub in candidate.publications],
        "applications": [{
            "job_id": app.job_id,
            "position_applied": app.position_applied,
            "admin_department": app.admin_department,
            "current_status": app.current_status,
            "submitted_at": app.submitted_at.isoformat() if app.submitted_at else None
        } for app in candidate.applications]
    }

    # Fetch stored AI evaluation from CandidateResumePayload
    import json
    payload = db.query(models.CandidateResumePayload).filter(
        models.CandidateResumePayload.candidate_id == candidate_id
    ).first()
    
    if payload and payload.ai_evaluation_json:
        try:
            res_profile["ai_evaluation"] = json.loads(payload.ai_evaluation_json)
        except json.JSONDecodeError:
            res_profile["ai_evaluation"] = None
    else:
        res_profile["ai_evaluation"] = None

    return res_profile


@app.get("/api/v1/hr/candidates/lookup", dependencies=[Depends(get_current_admin)])
def lookup_candidates(q: str, db: Session = Depends(get_db)):
    if not q or not q.strip():
        return []
        
    query_str = q.strip().lower()
    
    # Query candidates matching id (UUID/prefix), email, full_name, or mobile_no
    candidates = db.query(models.CandidateMetadata).options(
        joinedload(models.CandidateMetadata.higher_education),
        joinedload(models.CandidateMetadata.work_experiences),
        joinedload(models.CandidateMetadata.applications).joinedload(models.ApplicationTracking.job)
    ).filter(
        or_(
            func.lower(models.CandidateMetadata.id).like(f"%{query_str}%"),
            func.lower(models.CandidateMetadata.email).like(f"%{query_str}%"),
            func.lower(models.CandidateMetadata.full_name).like(f"%{query_str}%"),
            func.lower(models.CandidateMetadata.mobile_no).like(f"%{query_str}%")
        )
    ).limit(20).all()
    
    # If no direct metadata match, search by ApplicationTracking.id
    if not candidates:
        app_tracks = db.query(models.ApplicationTracking).filter(
            func.lower(models.ApplicationTracking.id).like(f"%{query_str}%")
        ).limit(20).all()
        cand_ids = [a.candidate_id for a in app_tracks if a.candidate_id]
        if cand_ids:
            candidates = db.query(models.CandidateMetadata).options(
                joinedload(models.CandidateMetadata.higher_education),
                joinedload(models.CandidateMetadata.work_experiences),
                joinedload(models.CandidateMetadata.applications).joinedload(models.ApplicationTracking.job)
            ).filter(models.CandidateMetadata.id.in_(cand_ids)).all()

    results = []
    for c in candidates:
        latest_app = c.applications[-1] if (c.applications and len(c.applications) > 0) else None
        job_obj = latest_app.job if (latest_app and hasattr(latest_app, 'job') and latest_app.job) else None
        
        pos_applied = (latest_app.position_applied if latest_app else None) or "Candidate"
        job_title = (job_obj.title if job_obj else pos_applied) or pos_applied
        
        submitted_at_str = ""
        if latest_app and latest_app.submitted_at:
            submitted_at_str = latest_app.submitted_at.strftime("%d-%m-%Y %H:%M IST") if hasattr(latest_app.submitted_at, 'strftime') else str(latest_app.submitted_at)[:16]
            
        top_edu = ""
        hedus = c.higher_education or []
        phds = [e for e in hedus if getattr(e, 'level', '') == 'phd']
        pgs = [e for e in hedus if getattr(e, 'level', '') == 'postgrad']
        ugs = [e for e in hedus if getattr(e, 'level', '') == 'undergrad']
        
        if phds:
            top_edu = f"Ph.D. ({phds[0].university or 'University'})"
        elif pgs:
            top_edu = f"{pgs[0].degree_name or 'Masters'} ({pgs[0].university or 'University'})"
        elif ugs:
            top_edu = f"{ugs[0].degree_name or 'Bachelors'} ({ugs[0].university or 'University'})"

        results.append({
            "candidate_id": c.id,
            "application_id": latest_app.id if latest_app else c.id,
            "full_name": c.full_name,
            "email": c.email,
            "mobile_no": f"{c.country_code or '+91'} {c.mobile_no or ''}".strip(),
            "city_state": f"{c.city or ''}, {c.state or ''}".strip(" ,") if not c.is_international_address else (c.international_address or "International"),
            "dob_age": f"{c.dob.strftime('%d-%m-%Y') if (c.dob and hasattr(c.dob, 'strftime')) else str(c.dob or '')[:10]} (Age: {c.age or 'N/A'})" if c.dob else "",
            "position_applied": pos_applied,
            "job_title": job_title,
            "admin_department": latest_app.admin_department if latest_app else "",
            "job_id": latest_app.job_id if latest_app else None,
            "submitted_at": submitted_at_str,
            "total_exp": f"{c.years_of_experience or 0.0} yrs",
            "top_edu": top_edu,
            "worked_at_ris": getattr(c, 'worked_at_ris', False) or False,
            "ris_designation": getattr(c, 'ris_designation', None)
        })
        
    return results


@app.post("/api/v1/candidates/{candidate_id}/ai_evaluate", dependencies=[Depends(get_current_admin)])
async def ai_evaluate_candidate(candidate_id: str, job_id: Optional[str] = None, db: Session = Depends(get_db)):
    candidate = db.query(models.CandidateMetadata).options(
        joinedload(models.CandidateMetadata.schooling),
        joinedload(models.CandidateMetadata.higher_education),
        joinedload(models.CandidateMetadata.publications),
        joinedload(models.CandidateMetadata.work_experiences),
        joinedload(models.CandidateMetadata.links_about),
        joinedload(models.CandidateMetadata.resume_payload),
        joinedload(models.CandidateMetadata.applications).joinedload(models.ApplicationTracking.job)
    ).filter(models.CandidateMetadata.id == candidate_id).first()

    if not candidate:
        raise HTTPException(status_code=404, detail=f"Candidate {candidate_id} not found")

    job = None
    if job_id:
        job = db.query(models.JobPosting).filter(models.JobPosting.id == job_id).first()
    elif candidate.applications and candidate.applications[0].job:
        job = candidate.applications[0].job

    job_title = job.title if job else "General Policy & Research Specialist"
    job_requirements = f"Title: {job_title}\nDescription: {job.description if job else 'Research & Policy analysis'}\nRequirements: {job.requirements if job else 'PhD/Masters, Publications, Relevant Policy experience'}"

    from services.ai_evaluator import evaluate_candidate_qualitative

    candidate_dict = {
        "full_name": candidate.full_name,
        "years_of_experience": candidate.years_of_experience,
        "degrees": [f"{e.level.upper()}: {e.degree_name} ({e.university})" for e in candidate.higher_education],
        "publications": [f"{p.pub_type.upper()}: {p.title}" for p in candidate.publications],
        "work_experiences": [f"{w.role} at {w.company_name}" for w in candidate.work_experiences],
        "sop": candidate.links_about.sop if candidate.links_about else "",
        "about": candidate.links_about.about if candidate.links_about else ""
    }

    import json
    try:
        eval_result = evaluate_candidate_qualitative(job_title, job_requirements, candidate_dict)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    eval_result["candidate_id"] = candidate_id

    # 💾 Save / Upsert AI Evaluation JSON in PostgreSQL database (candidate_resume_payload)
    payload = db.query(models.CandidateResumePayload).filter(
        models.CandidateResumePayload.candidate_id == candidate_id
    ).first()
    if not payload:
        payload = models.CandidateResumePayload(
            candidate_id=candidate_id,
            ai_evaluation_json=json.dumps(eval_result)
        )
        db.add(payload)
    else:
        payload.ai_evaluation_json = json.dumps(eval_result)
    
    db.commit()
    return eval_result


@app.get("/api/v1/applications/{candidate_id}/executive_dossier/download", dependencies=[Depends(get_current_admin)])
def download_executive_dossier(candidate_id: str, preview: bool = False, db: Session = Depends(get_db)):
    from fastapi.responses import Response
    import json
    try:
        candidate = db.query(models.CandidateMetadata).options(
            joinedload(models.CandidateMetadata.higher_education),
            joinedload(models.CandidateMetadata.publications),
            joinedload(models.CandidateMetadata.work_experiences),
            joinedload(models.CandidateMetadata.links_about),
            joinedload(models.CandidateMetadata.resume_payload),
            joinedload(models.CandidateMetadata.applications).joinedload(models.ApplicationTracking.job)
        ).filter(models.CandidateMetadata.id == candidate_id).first()

        if not candidate:
            raise HTTPException(status_code=404, detail="Candidate not found")

        payload = candidate.resume_payload
        ai_eval_data = None

        if payload and payload.ai_evaluation_json:
            try:
                ai_eval_data = json.loads(payload.ai_evaluation_json)
            except Exception:
                pass

        if not ai_eval_data:
            from services.ai_evaluator import evaluate_candidate_qualitative
            c_job = candidate.applications[0].job if (candidate.applications and candidate.applications[0].job) else None
            j_title = c_job.title if c_job else "General Research Specialist"
            j_reqs = c_job.requirements if c_job else "PhD/Masters, Policy Analysis"
            cand_dict = {
                "full_name": candidate.full_name or "Applicant",
                "years_of_experience": candidate.years_of_experience or 0,
                "degrees": [f"{e.level.upper() if e.level else 'DEGREE'}: {e.degree_name} ({e.university})" for e in (candidate.higher_education or [])],
                "publications": [f"{p.pub_type.upper() if p.pub_type else 'PUB'}: {p.title}" for p in (candidate.publications or [])],
                "work_experiences": [f"{w.role} at {w.company_name}" for w in (candidate.work_experiences or [])],
                "sop": candidate.links_about.sop if (candidate.links_about and candidate.links_about.sop) else "",
                "about": candidate.links_about.about if (candidate.links_about and candidate.links_about.about) else ""
            }
            ai_eval_data = evaluate_candidate_qualitative(j_title, j_reqs, cand_dict)
            if not payload:
                payload = models.CandidateResumePayload(candidate_id=candidate_id, ai_evaluation_json=json.dumps(ai_eval_data))
                db.add(payload)
            else:
                payload.ai_evaluation_json = json.dumps(ai_eval_data)
            db.commit()

        cand_profile_dict = {
            "full_name": candidate.full_name or "Applicant",
            "email": candidate.email or "N/A",
            "mobile_no": candidate.mobile_no or "",
            "country_code": candidate.country_code or "",
            "city": candidate.city or "",
            "state": candidate.state or "",
            "years_of_experience": candidate.years_of_experience or 0,
            "graduation": [{"degree_name": g.degree_name, "university": g.university} for g in (candidate.higher_education or [])],
            "sop": candidate.links_about.sop if (candidate.links_about and candidate.links_about.sop) else ""
        }

        from services.pdf_generator import generate_executive_dossier_pdf
        pdf_bytes = generate_executive_dossier_pdf(cand_profile_dict, ai_eval_data)

        disposition = "inline" if preview else "attachment"
        clean_name = (candidate.full_name or "Applicant").replace(" ", "_")
        filename = f"Executive_Dossier_{clean_name}.pdf"

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'{disposition}; filename="{filename}"'
            }
        )
    except Exception as route_err:
        print(f"❌ [Executive Dossier Download Route Error]: {route_err}")
        raise HTTPException(status_code=500, detail=f"Dossier PDF Generation Error: {str(route_err)}")


@app.get("/api/v1/applications/{candidate_id}/resume/download", dependencies=[Depends(get_current_admin)])
def download_resume(candidate_id: str, preview: bool = False, db: Session = Depends(get_db)):
    from fastapi.responses import RedirectResponse, FileResponse
    payload = db.query(models.CandidateResumePayload).filter(
        models.CandidateResumePayload.candidate_id == candidate_id
    ).first()
    
    if not payload or not payload.resume_path:
        raise HTTPException(status_code=404, detail="Resume record not found for candidate")

    filename = os.path.basename(payload.resume_path)
    S3_BUCKET = os.getenv("S3_BUCKET_NAME")
    aws_access_key = os.getenv("AWS_ACCESS_KEY_ID")
    aws_secret_key = os.getenv("AWS_SECRET_ACCESS_KEY")
    aws_region = os.getenv("AWS_REGION", "ap-south-1")

    if S3_BUCKET:
        try:
            import boto3
            from botocore.client import Config
            s3_kwargs = {"region_name": aws_region, "config": Config(signature_version='s3v4')}
            if aws_access_key and aws_secret_key:
                s3_kwargs["aws_access_key_id"] = aws_access_key
                s3_kwargs["aws_secret_access_key"] = aws_secret_key

            s3_client = boto3.client('s3', **s3_kwargs)
            disposition = "inline" if preview else "attachment"
            
            presigned_url = s3_client.generate_presigned_url(
                'get_object',
                Params={
                    'Bucket': S3_BUCKET,
                    'Key': payload.resume_path,
                    'ResponseContentDisposition': f'{disposition}; filename="{filename}"'
                },
                ExpiresIn=900
            )
            return RedirectResponse(url=presigned_url)
        except Exception as e:
            print(f"[S3 Download Error] Failed to generate S3 pre-signed URL: {e}")

    if payload.resume_path and os.path.exists(payload.resume_path):
        headers = {"Content-Disposition": f'{"inline" if preview else "attachment"}; filename="{filename}"'}
        return FileResponse(path=payload.resume_path, headers=headers, media_type="application/pdf")

    raise HTTPException(status_code=404, detail="CV file unavailable")


def process_and_save_resume(db: Session, candidate_id: str, file_bytes: bytes, filename: str, job_id: Optional[str] = None):
    candidate = db.query(models.CandidateMetadata).filter(models.CandidateMetadata.id == candidate_id).first()
    if not candidate:
        raise ValueError(f"Candidate {candidate_id} not found")

    # ── Job Categorization Folder Naming ──
    job_folder_name = "general_applications"
    app_rec = None
    if job_id:
        app_rec = db.query(models.ApplicationTracking).filter_by(candidate_id=candidate_id, job_id=job_id).first()
    if not app_rec:
        app_rec = db.query(models.ApplicationTracking).filter_by(candidate_id=candidate_id).order_by(models.ApplicationTracking.submitted_at.desc()).first()

    if app_rec:
        raw_title = "general_applications"
        if app_rec.job and app_rec.job.title:
            raw_title = app_rec.job.title
        elif app_rec.position_applied:
            raw_title = app_rec.position_applied
            
        import re
        job_folder_name = re.sub(r'[^a-zA-Z0-9_-]', '_', raw_title).strip('_')
        job_folder_name = re.sub(r'_+', '_', job_folder_name)

    content_type = get_resume_media_type(filename)
    S3_BUCKET = os.getenv("S3_BUCKET_NAME")
    aws_access_key = os.getenv("AWS_ACCESS_KEY_ID")
    aws_secret_key = os.getenv("AWS_SECRET_ACCESS_KEY")
    aws_region = os.getenv("AWS_REGION", "ap-south-1")
    
    # Categorized S3 Path Structure: jobs/{job_folder}/resumes/{candidate_id}_{filename}
    s3_key = f"jobs/{job_folder_name}/resumes/{candidate_id}_{filename}"

    if not S3_BUCKET:
        raise ValueError("AWS S3 bucket is not configured. Please set S3_BUCKET_NAME in environment.")

    # ── Direct AWS S3 Categorized Upload ──
    try:
        import boto3
        from botocore.client import Config
        s3_kwargs = {"region_name": aws_region, "config": Config(signature_version='s3v4')}
        if aws_access_key and aws_secret_key:
            s3_kwargs["aws_access_key_id"] = aws_access_key
            s3_kwargs["aws_secret_access_key"] = aws_secret_key

        s3_client = boto3.client('s3', **s3_kwargs)
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file_bytes,
            ContentType=content_type
        )
        print(f"✅ [S3 Categorized Upload] Saved under job folder: '{s3_key}'")
    except Exception as e:
        print(f"❌ [S3 Upload Error] Failed to upload resume to S3: {e}")
        raise ValueError(f"S3 Upload Failed: {str(e)}")

    # Store ONLY S3 path reference in PostgreSQL
    payload = db.query(models.CandidateResumePayload).filter(models.CandidateResumePayload.candidate_id == candidate_id).first()
    if payload:
        payload.resume_path = s3_key
        payload.pdf_blob = None
    else:
        payload = models.CandidateResumePayload(
            candidate_id=candidate_id, 
            resume_path=s3_key, 
            pdf_blob=None
        )
        db.add(payload)

    db.commit()
    return s3_key, s3_key

@app.post("/api/v1/applications/{candidate_id}/resume")
async def upload_resume(candidate_id: str, job_id: Optional[str] = None, file: UploadFile = File(...), db: Session = Depends(get_db)):
    try:
        content = await file.read()
        validate_resume_upload(file.filename, file.content_type, len(content))
        saved_path, file_path = process_and_save_resume(db, candidate_id, content, file.filename, job_id=job_id)
        return {"status": "success", "resume_path": saved_path}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class CandidateFilter(BaseModel):

    states: Optional[List[str]] = None
    nationalities: Optional[List[str]] = None
    genders: Optional[List[str]] = None
    ug_uni: Optional[str] = None
    min_ug_score: Optional[float] = None
    pg_uni: Optional[str] = None
    pg_min_score: Optional[float] = None
    phd_uni: Optional[str] = None
    phd_thesis: Optional[str] = None
    phd_min_score: Optional[float] = None
    min_experience_years: Optional[float] = None
    worked_at_ris: Optional[bool] = None
    min_papers: Optional[int] = None
    min_books: Optional[int] = None
    min_chapters: Optional[int] = None
    min_reports: Optional[int] = None
    min_policy_briefs: Optional[int] = None
    min_age: Optional[int] = None
    max_age: Optional[int] = None
    min_x_score: Optional[float] = None
    min_xii_score: Optional[float] = None
    role_keyword: Optional[str] = None
    company_keyword: Optional[str] = None
    # Score type awareness for education and schooling filters
    ug_score_type: Optional[str] = None    # 'Percentage' or 'CGPA'
    pg_score_type: Optional[str] = None
    phd_score_type: Optional[str] = None
    phd_domain: Optional[str] = None
    x_score_type: Optional[str] = None
    xii_score_type: Optional[str] = None


@app.get("/api/v1/universities")
def get_universities():
    from utils.scoring import UNIVERSITIES_DB
    return sorted(list(UNIVERSITIES_DB.keys()))


class ExportRequest(BaseModel):
    filters: CandidateFilter
    format: str # 'csv' or 'xlsx'
    columns: List[str]
    report_type: Optional[str] = 'detailed'


@app.post("/api/v1/jobs/{job_id}/candidates/export", dependencies=[Depends(get_current_admin)])
def export_job_candidates(job_id: str, req: ExportRequest, db: Session = Depends(get_db)):
    from io import StringIO, BytesIO
    import csv
    from fastapi.responses import StreamingResponse
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

    try:
        # 1. Fetch filtered candidates using the existing engine
        candidates_data = filter_job_candidates(job_id, req.filters, db)
        if not candidates_data:
            # Return empty response with headers
            if req.format == 'csv':
                return StreamingResponse(iter(["No candidates found"]), media_type="text/csv")
            else:
                output = BytesIO()
                wb = Workbook()
                wb.save(output)
                output.seek(0)
                return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        def format_score(value, score_type):
            if value is None or value == "":
                return ""
            try:
                val_num = float(value)
                val_str = f"{int(val_num)}" if val_num == int(val_num) else f"{val_num}"
            except Exception:
                val_str = str(value)

            stype = str(score_type or "").strip()
            stype_lower = stype.lower()
            
            if "percent" in stype_lower:
                return f"{val_str}%"
            elif "4" in stype:
                return f"{val_str} CGPA (Out of 4)"
            elif "10" in stype:
                return f"{val_str} CGPA (Out of 10)"
            elif "cgpa" in stype_lower:
                return f"{val_str} CGPA"
            else:
                return f"{val_str} {stype}".strip()

        def format_schooling_score(schooling, level):
            if not schooling:
                return ""
            if level == 'x':
                val = getattr(schooling, 'class_x_score_value', None)
                stype = getattr(schooling, 'class_x_score_type', '')
            else:
                val = getattr(schooling, 'class_xii_score_value', None)
                stype = getattr(schooling, 'class_xii_score_type', '')
                
            return format_score(val, stype)

        def calculate_age(dob):
            if not dob:
                return ""
            try:
                import datetime
                if isinstance(dob, str):
                    dob_date = datetime.date.fromisoformat(dob[:10])
                else:
                    dob_date = dob
                today = datetime.date.today()
                return today.year - dob_date.year - ((today.month, today.day) < (dob_date.month, dob_date.day))
            except Exception:
                return ""

        def set_cell_border(cell):
            try:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/></w:tcBorders>' % nsdecls('w'))
                tcPr.append(tcBorders)
            except Exception:
                pass

        def generate_docx_committee_brief(job, candidates):
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.section import WD_ORIENT
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_TABLE_ALIGNMENT
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                from io import BytesIO
                import datetime

                doc = Document()
                
                # Set Landscape Orientation A4
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
                section.top_margin = Inches(0.4)
                section.bottom_margin = Inches(0.4)
                section.left_margin = Inches(0.4)
                section.right_margin = Inches(0.4)
                
                # Document Header Title
                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                job_title_text = job.title if job and hasattr(job, 'title') and job.title else "Position"
                title_run = title_p.add_run(f"Brief of Applicants for {job_title_text}")
                title_run.font.name = 'Times New Roman'
                title_run.font.size = Pt(14)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0, 33, 71)
                title_p.paragraph_format.space_after = Pt(12)
                
                # 7-Column Table
                headers = [
                    "S. No.", 
                    "Basic Info", 
                    "X / XII", 
                    "Higher Education", 
                    "Work Experience", 
                    "Last Salary", 
                    "Publications"
                ]
                
                col_widths = [
                    Inches(0.5),  # S. No.
                    Inches(2.1),  # Basic Info
                    Inches(1.5),  # X / XII
                    Inches(2.5),  # Higher Education
                    Inches(2.5),  # Work Experience
                    Inches(0.9),  # Last Salary
                    Inches(1.3)   # Publications
                ]
                
                table = doc.add_table(rows=1, cols=7)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                try:
                    trPr = table.rows[0]._tr.get_or_add_trPr()
                    trPr.append(parse_xml(r'<w:tblHeader %s/>' % nsdecls('w')))
                except Exception:
                    pass
                
                hdr_cells = table.rows[0].cells
                for i, header_text in enumerate(headers):
                    hdr_cells[i].width = col_widths[i]
                    p = hdr_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(header_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    try:
                        shd = parse_xml(r'<w:shd %s w:fill="002147"/>' % nsdecls('w'))
                        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
                    except Exception:
                        pass
                    set_cell_border(hdr_cells[i])

                for s_idx, cand in enumerate(candidates, 1):
                    try:
                        row = table.add_row()
                        row_cells = row.cells
                        
                        try:
                            trPr = row._tr.get_or_add_trPr()
                            trPr.append(parse_xml(r'<w:cantSplit %s/>' % nsdecls('w')))
                        except Exception:
                            pass
                        
                        # Col 0: S. No.
                        p0 = row_cells[0].paragraphs[0]
                        r0 = p0.add_run(f"{s_idx}.")
                        r0.font.name = 'Times New Roman'
                        r0.font.size = Pt(9)
                        r0.font.bold = True
                        
                        # Col 1: Basic Info
                        full_name = cand.full_name or ""
                        
                        if getattr(cand, 'is_international_address', False) and getattr(cand, 'international_address', None):
                            addr_str = cand.international_address
                        else:
                            parts = [p for p in [getattr(cand, 'city', None), getattr(cand, 'state', None), getattr(cand, 'pincode', None)] if p]
                            addr_str = ", ".join(parts) if parts else "N/A"
                            
                        mobile_str = f"{cand.country_code or '+91'} {cand.mobile_no or 'N/A'}"
                        email_str = cand.email or 'N/A'
                        
                        dob_str = ""
                        age_str = ""
                        if getattr(cand, 'dob', None):
                            dob_str = cand.dob.strftime("%d-%m-%Y") if hasattr(cand.dob, 'strftime') else str(cand.dob)[:10]
                            today = datetime.date.today()
                            if hasattr(cand.dob, 'year'):
                                age = today.year - cand.dob.year - ((today.month, today.day) < (cand.dob.month, cand.dob.day))
                                age_str = f" (Age: {age} yrs)"
                            elif getattr(cand, 'age', None):
                                age_str = f" (Age: {cand.age} yrs)"
                        dob_age_line = f"DOB: {dob_str}{age_str}" if dob_str else ""

                        p1 = row_cells[1].paragraphs[0]
                        p1.paragraph_format.space_after = Pt(2)
                        r1_name = p1.add_run(f"{full_name}\n")
                        r1_name.font.name = 'Times New Roman'
                        r1_name.font.size = Pt(9)
                        r1_name.font.bold = True
                        r1_name.font.color.rgb = RGBColor(0, 33, 71)
                        
                        r1_details = p1.add_run(f"{addr_str}\n{mobile_str}\n{email_str}\n{dob_age_line}".strip())
                        r1_details.font.name = 'Times New Roman'
                        r1_details.font.size = Pt(8.5)
                        
                        # Col 2: X / XII
                        sch = getattr(cand, 'schooling', None)
                        p2 = row_cells[2].paragraphs[0]
                        p2.paragraph_format.space_after = Pt(2)
                        
                        x_xii_lines = []
                        if sch:
                            xii_sch = getattr(sch, 'class_xii_school', '') or ''
                            xii_brd = getattr(sch, 'class_xii_board', '') or ''
                            xii_yr = getattr(sch, 'class_xii_year', '') or ''
                            xii_val = getattr(sch, 'class_xii_score_value', 0.0)
                            xii_stype = getattr(sch, 'class_xii_score_type', '')
                            xii_score = format_score(xii_val, xii_stype) if xii_val else ""
                            if xii_sch or xii_brd or xii_yr:
                                x_xii_lines.append(f"• XII: {xii_sch}, {xii_brd}, {xii_yr} ({xii_score})".strip(" ,()"))

                            x_sch = getattr(sch, 'class_x_school', '') or ''
                            x_brd = getattr(sch, 'class_x_board', '') or ''
                            x_yr = getattr(sch, 'class_x_year', '') or ''
                            x_val = getattr(sch, 'class_x_score_value', 0.0)
                            x_stype = getattr(sch, 'class_x_score_type', '')
                            x_score = format_score(x_val, x_stype) if x_val else ""
                            if x_sch or x_brd or x_yr:
                                x_xii_lines.append(f"• X: {x_sch}, {x_brd}, {x_yr} ({x_score})".strip(" ,()"))
                            
                        r2 = p2.add_run("\n".join(x_xii_lines) if x_xii_lines else "N/A")
                        r2.font.name = 'Times New Roman'
                        r2.font.size = Pt(8.5)

                        # Col 3: Higher Education (Diploma, Bachelors, Masters, Ph.D.)
                        hedus = getattr(cand, 'higher_education', []) or []
                        p3 = row_cells[3].paragraphs[0]
                        p3.paragraph_format.space_after = Pt(2)
                        
                        hedu_lines = []
                        for h in hedus:
                            degree = getattr(h, 'degree_name', None) or getattr(h, 'degree_type', None) or 'Degree'
                            inst = getattr(h, 'university', '') or ''
                            
                            is_p = getattr(h, 'is_pursuing', False)
                            g_yr = getattr(h, 'grad_year', None)
                            if is_p and g_yr:
                                year = f"Pursuing - Exp. {g_yr}"
                            elif is_p:
                                year = "Pursuing"
                            elif g_yr:
                                year = str(g_yr)
                            else:
                                year = ""
                            
                            val = getattr(h, 'score_value', None)
                            stype = getattr(h, 'score_type', '')
                            marks = format_score(val, stype) if val else ""
                            
                            line = f"• {degree} - {inst} ({year}) {marks}".strip()
                            if getattr(h, 'level', '') == 'phd' or getattr(h, 'degree_type', '') == 'PhD':
                                if getattr(h, 'phd_domain', None):
                                    line += f"\n  Domain: {h.phd_domain}"
                                if getattr(h, 'thesis_title', None):
                                    line += f"\n  Thesis: {h.thesis_title}"
                            hedu_lines.append(line)
                            
                        r3 = p3.add_run("\n".join(hedu_lines) if hedu_lines else "N/A")
                        r3.font.name = 'Times New Roman'
                        r3.font.size = Pt(8.5)

                        # Col 4: Work Experience + RIS Experience Tag
                        works = getattr(cand, 'work_experiences', []) or []
                        p4 = row_cells[4].paragraphs[0]
                        p4.paragraph_format.space_after = Pt(2)
                        
                        total_months = 0
                        for w in works:
                            s = getattr(w, 'start_date', None)
                            e = getattr(w, 'end_date', None) or datetime.date.today()
                            if s and e:
                                total_months += (e.year - s.year) * 12 + (e.month - s.month)
                        
                        yrs = total_months // 12
                        mos = total_months % 12
                        exp_summary = f"Total Exp: {yrs} yrs {mos} mos\n" if total_months > 0 else "Total Exp: 0 yrs\n"
                        
                        r4_head = p4.add_run(exp_summary)
                        r4_head.font.name = 'Times New Roman'
                        r4_head.font.size = Pt(8.5)
                        r4_head.font.bold = True

                        # Check if candidate worked at RIS before
                        if getattr(cand, 'worked_at_ris', False):
                            ris_desig = getattr(cand, 'ris_designation', None) or "Prior Role"
                            ris_dates = ""
                            s_d = getattr(cand, 'ris_start_date', None)
                            e_d = getattr(cand, 'ris_end_date', None)
                            if s_d:
                                s_str = s_d.strftime("%b %Y") if hasattr(s_d, 'strftime') else str(s_d)[:7]
                                e_str = "Present" if getattr(cand, 'ris_is_current', False) else (e_d.strftime("%b %Y") if (e_d and hasattr(e_d, 'strftime')) else (str(e_d)[:7] if e_d else ''))
                                ris_dates = f" ({s_str} – {e_str})"
                            
                            r4_ris = p4.add_run(f"★ Prior RIS Experience: {ris_desig}{ris_dates}\n")
                            r4_ris.font.name = 'Times New Roman'
                            r4_ris.font.size = Pt(8.5)
                            r4_ris.font.bold = True
                            r4_ris.font.color.rgb = RGBColor(198, 40, 40)
                            
                        work_lines = []
                        for w in sorted(works, key=lambda x: getattr(x, 'start_date', None) or datetime.date(1900, 1, 1), reverse=True):
                            role = getattr(w, 'role', None) or getattr(w, 'designation', None) or 'Role'
                            company = getattr(w, 'company_name', None) or ''
                            s_d = getattr(w, 'start_date', None)
                            e_d = getattr(w, 'end_date', None)
                            s_str = s_d.strftime("%b %Y") if (s_d and hasattr(s_d, 'strftime')) else ''
                            e_str = "Present" if getattr(w, 'is_current', False) else (e_d.strftime("%b %Y") if (e_d and hasattr(e_d, 'strftime')) else '')
                            dates_str = f" ({s_str} – {e_str})" if s_str else ''
                            work_lines.append(f"• {role} at {company}{dates_str}")
                            
                        r4_works = p4.add_run("\n".join(work_lines) if work_lines else "None")
                        r4_works.font.name = 'Times New Roman'
                        r4_works.font.size = Pt(8.5)

                        # Col 5: Last Salary (LPA)
                        p5 = row_cells[5].paragraphs[0]
                        p5.paragraph_format.space_after = Pt(2)
                        last_sal = getattr(cand, 'last_salary', None)
                        sal_display = f"₹ {last_sal} LPA" if last_sal else "N/A"
                        r5 = p5.add_run(sal_display)
                        r5.font.name = 'Times New Roman'
                        r5.font.size = Pt(8.5)

                        # Col 6: Publications
                        pubs = getattr(cand, 'publications', []) or []
                        p6 = row_cells[6].paragraphs[0]
                        p6.paragraph_format.space_after = Pt(2)
                        
                        pub_counts = {}
                        for pub in pubs:
                            ptype = getattr(pub, 'pub_type', None) or getattr(pub, 'publication_type', None) or "Publication"
                            pub_counts[ptype] = pub_counts.get(ptype, 0) + 1
                            
                        pub_lines = [f"• {k.capitalize()}: {v}" for k, v in pub_counts.items() if v > 0]
                        r6 = p6.add_run("\n".join(pub_lines) if pub_lines else "None")
                        r6.font.name = 'Times New Roman'
                        r6.font.size = Pt(8.5)
                        
                        for i in range(7):
                            row_cells[i].width = col_widths[i]
                            set_cell_border(row_cells[i])
                    except Exception as _row_e:
                        print(f"Error rendering row {s_idx} for candidate {getattr(cand, 'id', 'unknown')}: {_row_e}")
                        continue

                docx_io = BytesIO()
                doc.save(docx_io)
                docx_io.seek(0)
                return docx_io
            except Exception as _doc_e:
                print(f"Error in generate_docx_committee_brief: {_doc_e}")
                # Fallback to minimal valid document
                from docx import Document
                from io import BytesIO
                err_doc = Document()
                err_doc.add_heading("Committee Brief Error", 0)
                err_doc.add_paragraph(f"Failed to generate brief: {str(_doc_e)}")
                err_io = BytesIO()
                err_doc.save(err_io)
                err_io.seek(0)
                return err_io

        # 2. Fetch job if applicable
        job_obj = db.query(models.JobPosting).filter(models.JobPosting.id == job_id).first() if job_id != 'all' else None

        # Handle Committee Brief (.docx) export
        if req.format == 'docx' or req.report_type == 'brief':
            candidate_objs = []
            for c in candidates_data:
                full_c = db.query(models.CandidateMetadata).filter(models.CandidateMetadata.id == c['id']).first()
                if full_c:
                    candidate_objs.append(full_c)
                    
            docx_stream = generate_docx_committee_brief(job_obj, candidate_objs)
            job_title_clean = (job_obj.title if job_obj else "Applicants").replace(" ", "_")
            filename = f"committee_brief_{job_title_clean}.docx"
            return StreamingResponse(
                docx_stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )

        # 3. Build Headers and Rows based on report type
        if req.report_type == 'standardized':
            headers = [
                "Full Name", "Date of Birth", "Age", "Email", "Mobile No", "Gender", "Nationality",
                "Is International Address?", "International Address", "City / State / Pin",
                "Position Applied", "Division / Department", "Current Status", "Submitted Date", "Source",
                "LinkedIn Link", "Worked for RIS Before?", "Class X Score", "Class X Year", "Class XII Score", "Class XII Year", 
                "Bachelors (UG)", "Bachelors Score", "Bachelors Year",
                "Masters (PG)", "Masters Score", "Masters Year",
                "Doctorate (PhD) Thesis Title", "PhD Main Domain", "Doctorate Year of Award",
                "Total Exp (Yrs)", "Latest Employment"
            ]
            
            rows_to_write = []
            for c in candidates_data:
                full_c = db.query(models.CandidateMetadata).filter(models.CandidateMetadata.id == c['id']).first()
                if not full_c: continue
                
                # Fetch qualifications
                ug = next((e for e in full_c.higher_education if e.level == 'undergrad'), None)
                pg = next((e for e in full_c.higher_education if e.level == 'postgrad'), None)
                phd = next((e for e in full_c.higher_education if e.level == 'phd'), None)
                
                # Get latest work experience (by date or highest order)
                latest_work = None
                if full_c.work_experiences:
                    latest_work = max(full_c.work_experiences, key=lambda x: x.start_date)
                
                ug_text = f"{ug.degree_name} ({ug.university})" if (ug and ug.university) else (ug.degree_name if ug else "")
                pg_text = f"{pg.degree_name} ({pg.university})" if (pg and pg.university) else (pg.degree_name if pg else "")
                phd_text = f"{phd.degree_name} ({phd.university})" if (phd and phd.university) else (phd.degree_name if phd else "")
                phd_domain_text = (phd.phd_domain if phd else "") or ""
                phd_year_text = (phd.grad_year if phd else "") or ""
                
                latest_work_text = f"{latest_work.role} ({latest_work.company_name})" if latest_work else ""
                
                latest_app = full_c.applications[-1] if getattr(full_c, 'applications', None) else None
                app_job = latest_app.job if (latest_app and hasattr(latest_app, 'job') and latest_app.job) else None

                pos_applied = (latest_app.position_applied if latest_app else None) or c.get('position_applied', '') or ""
                admin_dept = (latest_app.admin_department if latest_app else None) or c.get('admin_department', '') or ""
                current_stat = (latest_app.current_status if latest_app else None) or c.get('current_status', 'received')
                sub_date = (str(latest_app.submitted_at)[:10] if (latest_app and latest_app.submitted_at) else "") or (str(c.get('submitted_at', ''))[:10] if c.get('submitted_at') else "")

                # Derive Division / Department matching DB consistency
                if app_job and app_job.division:
                    div_dept = app_job.division
                elif job_obj and job_obj.division:
                    div_dept = job_obj.division
                elif pos_applied == 'Admin' and admin_dept:
                    div_dept = f"Admin - {admin_dept}"
                elif admin_dept:
                    div_dept = f"Admin - {admin_dept}"
                else:
                    div_dept = pos_applied

                is_intl = "Yes" if getattr(full_c, 'is_international_address', False) else "No"
                intl_addr = getattr(full_c, 'international_address', None) or ""
                
                if full_c.is_international_address or intl_addr:
                    location_str = intl_addr or "International Address"
                else:
                    parts = [p for p in [full_c.city, full_c.state, full_c.pincode] if p]
                    location_str = " - ".join(parts) if parts else ""

                ris_exp_str = f"Yes ({full_c.ris_designation or 'Worked at RIS'}: {str(full_c.ris_start_date or '')[:7]} to {'Present' if full_c.ris_is_current else str(full_c.ris_end_date or '')[:7]})" if getattr(full_c, 'worked_at_ris', False) else "No"

                row = [
                    full_c.full_name,
                    str(full_c.dob) if full_c.dob else "",
                    calculate_age(full_c.dob) if full_c.dob else (full_c.age or ""),
                    full_c.email,
                    f"{full_c.country_code or ''} {full_c.mobile_no}".strip(),
                    full_c.gender or "",
                    getattr(full_c, 'nationality', None) or "Indian",
                    is_intl,
                    intl_addr,
                    location_str,
                    pos_applied,
                    div_dept,
                    current_stat,
                    sub_date,
                    (full_c.links_about.how_heard if full_c.links_about else "") or "",
                    (full_c.links_about.linkedin if full_c.links_about else "") or "",
                    ris_exp_str,
                    format_schooling_score(full_c.schooling, "x"),
                    full_c.schooling.class_x_year if (full_c.schooling and full_c.schooling.class_x_year) else "",
                    format_schooling_score(full_c.schooling, "xii"),
                    full_c.schooling.class_xii_year if (full_c.schooling and full_c.schooling.class_xii_year) else "",
                    ug_text,
                    format_score(ug.score_value, ug.score_type) if ug else "",
                    ug.grad_year if ug else "",
                    pg_text,
                    format_score(pg.score_value, pg.score_type) if pg else "",
                    pg.grad_year if pg else "",
                    phd_text,
                    phd_domain_text,
                    phd_year_text,
                    full_c.years_of_experience or 0.0,
                    latest_work_text
                ]
                rows_to_write.append(row)
                
            merge_ranges = []
            candidate_groups = []
            single_line_rows = []
        else:
            # Detailed Report (Grouped Roster with Complete Input Fields, Booleans & Reorganized Publications)
            headers = [
                # Personal & Contact (1-13)
                "Full Name", "Email", "Country Code", "Mobile No", "Date of Birth", "Age", "Gender", 
                "Nationality", "Is International Address?", "International Address", "State", "City", "Pincode",
                # Application & Source (14-18)
                "Position Applied", "Division / Department", "Current Status", "Submitted Date", "Source (Where heard)",
                # Profiles & SOP (19-21)
                "Statement of Purpose (SOP)", "Google Scholar Link", "LinkedIn Link",
                # Boolean Indicator Flags (22-30)
                "Has Work Experience", "Currently Working", "Has Higher Education", "Currently Pursuing Degree", 
                "Has Doctorate (PhD)", "Has Master Degree (PG)", "Has Bachelor Degree (UG)", "Has Diploma", "Has Publications",
                # Schooling Class X & XII (31-38)
                "Class X School", "Class X Board", "Class X Score", "Class X Year",
                "Class XII School", "Class XII Board", "Class XII Score", "Class XII Year",
                # Graduation (39-43)
                "Graduation Univ", "Graduation Degree", "Graduation Score", "Graduation Year", "Graduation Pursuing",
                # Postgrad (44-48)
                "Postgrad Univ", "Postgrad Degree", "Postgrad Score", "Postgrad Year", "Postgrad Pursuing",
                # PhD (49-53)
                "PhD Univ", "PhD Main Domain", "PhD Thesis Title / Spec", "PhD Year of Award", "PhD Pursuing",
                # Diploma (54-58)
                "Diploma Institute", "Diploma Degree / Type", "Diploma Score", "Diploma Year", "Diploma Pursuing",
                # Experience & Salary (59-67)
                "Total Exp (Yrs)", "Last Salary (LPA)", "Worked for RIS Before?", "RIS Position / Role", "RIS Employment Period", "Work Organization", "Work Designation", "Work Start Date", "Work End Date",
                # Publications Reorganized Count -> Links (68-77)
                "Books & Book Chapters Count", "Books & Book Chapters Validation Links",
                "Peer-Reviewed Papers Count", "Peer-Reviewed Papers Validation Links",
                "Preprints & Chapters Count", "Preprints & Chapters Validation Links",
                "Research Reports Count", "Research Reports Validation Links",
                "Policy Briefs & Public Commentary Count", "Policy Briefs & Public Commentary Validation Links"
            ]
            
            rows_to_write = []
            merge_ranges = []
            candidate_groups = []
            single_line_rows = []
            current_r = 2 # Row 1 is headers
            
            import re

            def extract_pub_link_list(candidate_obj, match_types):
                links = []
                if candidate_obj.publications:
                    for p in candidate_obj.publications:
                        ptype = str(p.pub_type or "").lower().strip()
                        if ptype in match_types and p.title and p.title.strip():
                            raw = p.title.strip()
                            parts = [x.strip() for x in re.split(r'[\r\n;,]+', raw) if x.strip()]
                            links.extend(parts)
                return links

            for c in candidates_data:
                full_c = db.query(models.CandidateMetadata).filter(models.CandidateMetadata.id == c['id']).first()
                if not full_c: continue
                
                undergrads = [e for e in full_c.higher_education if e.level == 'undergrad']
                postgrads = [e for e in full_c.higher_education if e.level == 'postgrad']
                phds = [e for e in full_c.higher_education if e.level == 'phd']
                diplomas = [e for e in full_c.higher_education if e.level == 'diploma']
                works = sorted(full_c.work_experiences, key=lambda x: x.entry_order if x.entry_order else 1)

                book_links = extract_pub_link_list(full_c, ['book', 'books', 'book chapter', 'books & book chapters'])
                paper_links = extract_pub_link_list(full_c, ['paper', 'papers', 'journal', 'article', 'peer-reviewed journal papers'])
                chapter_links = extract_pub_link_list(full_c, ['chapter', 'chapters', 'preprint', 'working paper', 'working papers & preprints'])
                report_links = extract_pub_link_list(full_c, ['report', 'reports', 'research report', 'research reports & policy briefs'])
                brief_links = extract_pub_link_list(full_c, ['policy_brief', 'policy_briefs', 'brief', 'briefs', 'commentary', 'newspaper articles & public commentary'])

                max_rows = max(
                    len(undergrads), len(postgrads), len(phds), len(diplomas), len(works),
                    len(book_links), len(paper_links), len(chapter_links), len(report_links), len(brief_links), 1
                )
                
                candidate_groups.append((current_r, current_r + max_rows - 1))
                
                if max_rows == 1:
                    single_line_rows.append(current_r)
                
                if max_rows > 1:
                    # Merge candidate static metadata columns (Cols 1-38, 59-63, 68, 70, 72, 74, 76)
                    merge_cols = list(range(1, 39)) + [59, 60, 61, 62, 63, 68, 70, 72, 74, 76]
                    for col in merge_cols:
                        merge_ranges.append((current_r, current_r + max_rows - 1, col))

                latest_app = full_c.applications[-1] if getattr(full_c, 'applications', None) else None
                app_job = latest_app.job if (latest_app and hasattr(latest_app, 'job') and latest_app.job) else None

                pos_applied = (latest_app.position_applied if latest_app else None) or c.get('position_applied', '') or ""
                admin_dept = (latest_app.admin_department if latest_app else None) or c.get('admin_department', '') or ""
                current_stat = (latest_app.current_status if latest_app else None) or c.get('current_status', 'received')
                sub_date = (str(latest_app.submitted_at)[:10] if (latest_app and latest_app.submitted_at) else "") or (str(c.get('submitted_at', ''))[:10] if c.get('submitted_at') else "")

                # Derive Division / Department
                if app_job and app_job.division:
                    div_dept = app_job.division
                elif job_obj and job_obj.division:
                    div_dept = job_obj.division
                elif pos_applied == 'Admin' and admin_dept:
                    div_dept = f"Admin - {admin_dept}"
                elif admin_dept:
                    div_dept = f"Admin - {admin_dept}"
                else:
                    div_dept = pos_applied

                for i in range(max_rows):
                    row = [""] * len(headers)
                    if i == 0:
                        # Personal & Contact (0-12 index -> Cols 1-13)
                        row[0] = full_c.full_name
                        row[1] = full_c.email
                        row[2] = full_c.country_code or "+91"
                        row[3] = full_c.mobile_no
                        row[4] = str(full_c.dob) if full_c.dob else ""
                        row[5] = calculate_age(full_c.dob) if full_c.dob else (full_c.age or "")
                        row[6] = full_c.gender or ""
                        row[7] = getattr(full_c, 'nationality', None) or "Indian"
                        row[8] = "Yes" if getattr(full_c, 'is_international_address', False) else "No"
                        row[9] = getattr(full_c, 'international_address', None) or ""
                        row[10] = full_c.state or ""
                        row[11] = full_c.city or ""
                        row[12] = full_c.pincode or ""
                        
                        # Application & Source (13-17 index -> Cols 14-18)
                        row[13] = pos_applied
                        row[14] = div_dept
                        row[15] = current_stat
                        row[16] = sub_date
                        row[17] = (full_c.links_about.how_heard if full_c.links_about else "") or ""
                        
                        # Profiles & SOP (18-20 index -> Cols 19-21)
                        row[18] = (full_c.links_about.sop if full_c.links_about else "") or ""
                        row[19] = (full_c.links_about.google_scholar if full_c.links_about else "") or ""
                        row[20] = (full_c.links_about.linkedin if full_c.links_about else "") or ""

                        # Boolean Indicator Flags (21-29 index -> Cols 22-30)
                        row[21] = "Yes" if works else "No"
                        row[22] = "Yes" if any(w.is_current or not w.end_date for w in works) else "No"
                        row[23] = "Yes" if full_c.higher_education else "No"
                        row[24] = "Yes" if any(e.is_pursuing for e in full_c.higher_education) else "No"
                        row[25] = "Yes" if phds else "No"
                        row[26] = "Yes" if postgrads else "No"
                        row[27] = "Yes" if undergrads else "No"
                        row[28] = "Yes" if diplomas else "No"
                        row[29] = "Yes" if (full_c.publications or (full_c.links_about and (full_c.links_about.pub_books or full_c.links_about.pub_papers or full_c.links_about.pub_chapters or full_c.links_about.pub_reports or full_c.links_about.pub_policy_briefs))) else "No"

                        # Schooling Class X & XII (30-37 index -> Cols 31-38)
                        row[30] = full_c.schooling.class_x_school if full_c.schooling else ""
                        row[31] = full_c.schooling.class_x_board if full_c.schooling else ""
                        row[32] = format_schooling_score(full_c.schooling, "x")
                        row[33] = full_c.schooling.class_x_year if (full_c.schooling and full_c.schooling.class_x_year) else ""

                        row[34] = full_c.schooling.class_xii_school if full_c.schooling else ""
                        row[35] = full_c.schooling.class_xii_board if full_c.schooling else ""
                        row[36] = format_schooling_score(full_c.schooling, "xii")
                        row[37] = full_c.schooling.class_xii_year if (full_c.schooling and full_c.schooling.class_xii_year) else ""

                        # Experience & Salary Summary (58-62 index -> Cols 59-63)
                        row[58] = full_c.years_of_experience if full_c.years_of_experience is not None else 0.0
                        row[59] = full_c.last_salary if full_c.last_salary is not None else ""
                        row[60] = "Yes" if getattr(full_c, 'worked_at_ris', False) else "No"
                        row[61] = getattr(full_c, 'ris_designation', None) or ""
                        row[62] = f"{getattr(full_c, 'ris_start_date', '') or ''} to {'Present' if getattr(full_c, 'ris_is_current', False) else (getattr(full_c, 'ris_end_date', '') or '')}" if getattr(full_c, 'worked_at_ris', False) else ""

                        # Publications Summary Counts (67, 69, 71, 73, 75 index -> Cols 68, 70, 72, 74, 76)
                        row[67] = (full_c.links_about.pub_books if full_c.links_about else 0) or len(book_links)
                        row[69] = (full_c.links_about.pub_papers if full_c.links_about else 0) or len(paper_links)
                        row[71] = (full_c.links_about.pub_chapters if full_c.links_about else 0) or len(chapter_links)
                        row[73] = (full_c.links_about.pub_reports if full_c.links_about else 0) or len(report_links)
                        row[75] = (full_c.links_about.pub_policy_briefs if full_c.links_about else 0) or len(brief_links)

                    # Sub-Row i Specific Values (Unmerged Columns)
                    # Graduation details (38-42 index -> Cols 39-43)
                    if i < len(undergrads):
                        row[38] = undergrads[i].university or ""
                        row[39] = undergrads[i].degree_name or ""
                        row[40] = format_score(undergrads[i].score_value, undergrads[i].score_type) if undergrads[i].score_value else ""
                        row[41] = undergrads[i].grad_year or ""
                        row[42] = "Yes" if undergrads[i].is_pursuing else "No"
                    
                    # Postgrad details (43-47 index -> Cols 44-48)
                    if i < len(postgrads):
                        row[43] = postgrads[i].university or ""
                        row[44] = postgrads[i].degree_name or ""
                        row[45] = format_score(postgrads[i].score_value, postgrads[i].score_type) if postgrads[i].score_value else ""
                        row[46] = postgrads[i].grad_year or ""
                        row[47] = "Yes" if postgrads[i].is_pursuing else "No"
                    
                    # PhD details (48-52 index -> Cols 49-53)
                    if i < len(phds):
                        row[48] = phds[i].university or ""
                        row[49] = phds[i].phd_domain or ""
                        row[50] = phds[i].degree_name or ""
                        row[51] = phds[i].grad_year or ""
                        row[52] = "Yes" if phds[i].is_pursuing else "No"

                    # Diploma details (53-57 index -> Cols 54-58)
                    if i < len(diplomas):
                        row[53] = diplomas[i].university or ""
                        row[54] = diplomas[i].degree_name or ""
                        row[55] = format_score(diplomas[i].score_value, diplomas[i].score_type) if diplomas[i].score_value else ""
                        row[56] = diplomas[i].grad_year or ""
                        row[57] = "Yes" if diplomas[i].is_pursuing else "No"
                    
                    # General Work Experience details (63-66 index -> Cols 64-67)
                    if i < len(works):
                        row[63] = works[i].company_name or ""
                        row[64] = works[i].role or ""
                        row[65] = str(works[i].start_date) if works[i].start_date else ""
                        row[66] = str(works[i].end_date or "Present") if works[i].start_date else ""

                    # Publication Validation Links (Per Sub-Row i)
                    if i < len(book_links):
                        row[68] = book_links[i]
                    if i < len(paper_links):
                        row[70] = paper_links[i]
                    if i < len(chapter_links):
                        row[72] = chapter_links[i]
                    if i < len(report_links):
                        row[74] = report_links[i]
                    if i < len(brief_links):
                        row[76] = brief_links[i]

                    rows_to_write.append(row)
                    current_r += 1

        # 3. Generate Output
        if req.format == 'csv':
            si = StringIO()
            cw = csv.writer(si)
            cw.writerow(headers)
            cw.writerows(rows_to_write)
            si.seek(0)
            return StreamingResponse(
                iter([si.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=applicants_{job_id}.csv"}
            )
        else:
            # XLSX using openpyxl directly
            output = BytesIO()
            wb = Workbook()
            ws = wb.active
            ws.title = "Applicants"
            
            # Styling
            header_fill = PatternFill(start_color='1E3A8A', end_color='1E3A8A', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFF', size=11)
            
            divider_cols = [13, 18, 21, 30, 38, 43, 48, 53, 58, 63, 67, 69, 71, 73, 75, 77]

            # Write Headers
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                
                # Header vertical dividers for detailed view
                r_style = 'medium' if (req.report_type == 'detailed' and col_idx in divider_cols) else 'thin'
                r_color = '1E3A8A' if (req.report_type == 'detailed' and col_idx in divider_cols) else 'CBD5E1'
                
                cell.border = Border(
                    left=Side(style='thin', color='CBD5E1'),
                    right=Side(style=r_style, color=r_color),
                    top=Side(style='thin', color='CBD5E1'),
                    bottom=Side(style='medium', color='1E3A8A')
                )
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

            ws.row_dimensions[1].height = 40

            # Write Data
            for r_idx, row_data in enumerate(rows_to_write, 2):
                for c_idx, value in enumerate(row_data, 1):
                    cell = ws.cell(row=r_idx, column=c_idx, value=value)
                    
                    # Right border vertical dividers
                    r_style = 'medium' if (req.report_type == 'detailed' and c_idx in divider_cols) else 'thin'
                    r_color = '1E3A8A' if (req.report_type == 'detailed' and c_idx in divider_cols) else 'CBD5E1'
                    
                    cell.border = Border(
                        left=Side(style='thin', color='CBD5E1'),
                        right=Side(style=r_style, color=r_color),
                        top=Side(style='thin', color='CBD5E1'),
                        bottom=Side(style='thin', color='CBD5E1')
                    )
                    
                    # Alignments
                    if req.report_type == 'standardized':
                        if c_idx in [2, 3, 5, 6, 7, 8, 13, 14, 15, 17, 18, 20, 21, 23, 24, 26, 27, 29, 30]:
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                        else:
                            cell.alignment = Alignment(horizontal='left', vertical='center')
                    else:
                        # Detailed aligns
                        if c_idx in [3, 4, 5, 6, 7, 8, 9, 13, 16, 17, 22, 23, 24, 25, 26, 27, 28, 29, 30, 33, 34, 37, 38, 41, 42, 43, 46, 47, 48, 51, 52, 53, 56, 57, 58, 59, 60, 61, 63, 66, 67, 68, 70, 72, 74, 76]:
                            cell.alignment = Alignment(horizontal='center', vertical='top', wrap_text=True)
                        else:
                            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

            # Apply merges (only for detailed report type)
            if req.report_type == 'detailed':
                for start_r, end_r, col in merge_ranges:
                    ws.merge_cells(start_row=start_r, start_column=col, end_row=end_r, end_column=col)
                    h_align = 'center' if col in [3, 4, 5, 6, 7, 8, 9, 13, 16, 17, 22, 23, 24, 25, 26, 27, 28, 29, 30, 33, 34, 37, 38, 59, 60, 61, 63, 68, 70, 72, 74, 76] else 'left'
                    ws.cell(row=start_r, column=col).alignment = Alignment(vertical='top', horizontal=h_align, wrap_text=True)

                # Set bottom boundaries borders for groups (preserving vertical dividers)
                for start_r, end_r in candidate_groups:
                    for col in range(1, len(headers) + 1):
                        cell = ws.cell(row=end_r, column=col)
                        r_style = 'medium' if col in divider_cols else 'thin'
                        r_color = '1E3A8A' if col in divider_cols else 'CBD5E1'
                        cell.border = Border(
                            left=Side(style='thin', color='CBD5E1'),
                            right=Side(style=r_style, color=r_color),
                            top=cell.border.top or Side(style='thin', color='CBD5E1'),
                            bottom=Side(style='medium', color='1E3A8A')
                        )
                
                # Apply row heights for single line rows (26px height)
                for row_idx in single_line_rows:
                    ws.row_dimensions[row_idx].height = 26

            # Auto-fit column widths based on maximum content length
            from openpyxl.utils import get_column_letter
            for col in ws.columns:
                max_len = 0
                col_letter = get_column_letter(col[0].column)
                for cell in col:
                    if cell.value:
                        for line in str(cell.value).split('\n'):
                            max_len = max(max_len, len(str(line)))
                ws.column_dimensions[col_letter].width = min(max(max_len + 4, 16), 65)

            wb.save(output)
            output.seek(0)
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=applicants_{job_id}.xlsx"}
            )


    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Autocomplete Suggestion Endpoint ──
@app.get("/api/v1/jobs/{job_id}/suggest", dependencies=[Depends(get_current_admin)])
def suggest_tokens(job_id: str, field: str, q: str = "", db: Session = Depends(get_db)):
    """
    Returns substring-matched suggestions for filter fields.
    field can be: university, degree, company, role, paper, book, chapter
    """
    clean_id = str(job_id).strip()
    
    if not q or len(q) < 1:
        # Return top 15 most frequent tokens of this type
        results = db.query(models.TokenRegistry.token_value).filter(
            models.TokenRegistry.job_id == clean_id,
            models.TokenRegistry.token_type == field
        ).order_by(models.TokenRegistry.frequency.desc()).limit(15).all()
        return [r[0] for r in results]
    
    # Substring match (ILIKE '%q%')
    results = db.query(models.TokenRegistry.token_value).filter(
        models.TokenRegistry.job_id == clean_id,
        models.TokenRegistry.token_type == field,
        models.TokenRegistry.normalized.ilike(f"%{q.lower().strip()}%")
    ).order_by(models.TokenRegistry.frequency.desc()).limit(15).all()
    
    return [r[0] for r in results]

@app.post("/api/v1/jobs/{job_id}/candidates/filter", dependencies=[Depends(get_current_admin)])
def filter_job_candidates(job_id: str, filters: CandidateFilter, db: Session = Depends(get_db)):
    try:
        clean_job_id = str(job_id).strip()
        
        # 1. Base query for IDs joining application_tracking
        id_query = db.query(models.CandidateMetadata.id).join(
            models.ApplicationTracking, models.ApplicationTracking.candidate_id == models.CandidateMetadata.id
        )
        if clean_job_id not in ['all', 'all_jobs', '']:
            id_query = id_query.filter(models.ApplicationTracking.job_id == clean_job_id)

        # Apply personal filters
        if filters.states and len(filters.states) > 0:
            id_filters = [models.CandidateMetadata.state.ilike(f"%{s}%") for s in filters.states]
            id_query = id_query.filter(or_(*id_filters))
        
        if filters.nationalities and len(filters.nationalities) > 0:
            nat_filters = [models.CandidateMetadata.nationality.ilike(f"%{n}%") for n in filters.nationalities]
            id_query = id_query.filter(or_(*nat_filters))

        if filters.genders and len(filters.genders) > 0:
            id_query = id_query.filter(models.CandidateMetadata.gender.in_(filters.genders))

        if filters.min_experience_years is not None:
            id_query = id_query.filter(models.CandidateMetadata.years_of_experience >= float(filters.min_experience_years))

        if filters.min_age is not None:
            id_query = id_query.filter(models.CandidateMetadata.age >= filters.min_age)

        if filters.max_age is not None:
            id_query = id_query.filter(models.CandidateMetadata.age <= filters.max_age)

        def apply_score_type_filter(query_obj, db_column, target_score_type):
            if not target_score_type or str(target_score_type).strip() == "":
                return query_obj
            stype = str(target_score_type).strip().lower()
            if "percent" in stype:
                return query_obj.filter(db_column.ilike('%percent%'))
            elif "4" in stype:
                return query_obj.filter(or_(db_column.ilike('%4%'), db_column == 'CGPA (Out of 4)', db_column == 'CGPA_4'))
            elif "10" in stype:
                return query_obj.filter(or_(db_column.ilike('%10%'), db_column == 'CGPA (Out of 10)', db_column == 'CGPA_10'))
            elif "cgpa" in stype:
                return query_obj.filter(db_column.ilike('%cgpa%'))
            else:
                return query_obj.filter(db_column.ilike(f"%{target_score_type}%"))

        # Higher Education filters (UG/PG/PhD)
        if filters.ug_uni:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'undergrad',
                models.CandidateHigherEducation.university.ilike(f"%{filters.ug_uni}%")
            ).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))
        
        if filters.min_ug_score is not None:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'undergrad',
                models.CandidateHigherEducation.score_value >= float(filters.min_ug_score)
            )
            if filters.ug_score_type:
                sub = apply_score_type_filter(sub, models.CandidateHigherEducation.score_type, filters.ug_score_type)
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub.subquery()))

        if filters.pg_uni:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'postgrad',
                models.CandidateHigherEducation.university.ilike(f"%{filters.pg_uni}%")
            ).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))

        if filters.pg_min_score is not None:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'postgrad',
                models.CandidateHigherEducation.score_value >= float(filters.pg_min_score)
            )
            if filters.pg_score_type:
                sub = apply_score_type_filter(sub, models.CandidateHigherEducation.score_type, filters.pg_score_type)
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub.subquery()))

        if filters.phd_uni:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'phd',
                models.CandidateHigherEducation.university.ilike(f"%{filters.phd_uni}%")
            ).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))

        if filters.phd_thesis:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'phd',
                models.CandidateHigherEducation.degree_name.ilike(f"%{filters.phd_thesis}%") # PhD degree name = thesis title
            ).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))

        if filters.phd_domain:
            sub = db.query(models.CandidateHigherEducation.candidate_id).filter(
                models.CandidateHigherEducation.level == 'phd',
                models.CandidateHigherEducation.phd_domain.ilike(f"%{filters.phd_domain}%")
            ).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))

        # Academic Schooling Filters
        if filters.min_x_score is not None:
            sub = db.query(models.CandidateSchooling.candidate_id).filter(
                models.CandidateSchooling.class_x_score_value >= float(filters.min_x_score)
            )
            if filters.x_score_type:
                sub = apply_score_type_filter(sub, models.CandidateSchooling.class_x_score_type, filters.x_score_type)
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub.subquery()))

        if filters.min_xii_score is not None:
            sub = db.query(models.CandidateSchooling.candidate_id).filter(
                models.CandidateSchooling.class_xii_score_value >= float(filters.min_xii_score)
            )
            if filters.xii_score_type:
                sub = apply_score_type_filter(sub, models.CandidateSchooling.class_xii_score_type, filters.xii_score_type)
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub.subquery()))

        # Work Experience Filters
        if filters.role_keyword:
            sub = db.query(models.CandidateWorkExperience.candidate_id).filter(models.CandidateWorkExperience.role.ilike(f"%{filters.role_keyword}%")).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))

        if filters.company_keyword:
            sub = db.query(models.CandidateWorkExperience.candidate_id).filter(models.CandidateWorkExperience.company_name.ilike(f"%{filters.company_keyword}%")).subquery()
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub))

        if filters.worked_at_ris:
            id_query = id_query.filter(models.CandidateMetadata.worked_at_ris == True)

        # Publication Filters (now tracked in CandidateLinksAbout)
        if any(v is not None and v > 0 for v in [filters.min_papers, filters.min_books, filters.min_chapters, filters.min_reports, filters.min_policy_briefs]):
            sub = db.query(models.CandidateLinksAbout.candidate_id)
            if filters.min_papers and filters.min_papers > 0:
                sub = sub.filter(models.CandidateLinksAbout.pub_papers >= filters.min_papers)
            if filters.min_books and filters.min_books > 0:
                sub = sub.filter(models.CandidateLinksAbout.pub_books >= filters.min_books)
            if filters.min_chapters and filters.min_chapters > 0:
                sub = sub.filter(models.CandidateLinksAbout.pub_chapters >= filters.min_chapters)
            if filters.min_reports and filters.min_reports > 0:
                sub = sub.filter(models.CandidateLinksAbout.pub_reports >= filters.min_reports)
            if filters.min_policy_briefs and filters.min_policy_briefs > 0:
                sub = sub.filter(models.CandidateLinksAbout.pub_policy_briefs >= filters.min_policy_briefs)
            
            id_query = id_query.filter(models.CandidateMetadata.id.in_(sub.subquery()))

        # 2. Get the final list of matching IDs
        matching_ids = [r[0] for r in id_query.all()]

        if not matching_ids:
            return []

        # 3. Fetch full objects with joinedload only for those IDs
        candidates = db.query(models.CandidateMetadata).filter(
            models.CandidateMetadata.id.in_(matching_ids)
        ).options(
            joinedload(models.CandidateMetadata.higher_education),
            joinedload(models.CandidateMetadata.schooling),
            joinedload(models.CandidateMetadata.publications),
            joinedload(models.CandidateMetadata.work_experiences)
        ).all()

        # Fetch matching application trackers for status
        trackers = db.query(models.ApplicationTracking).filter(
            models.ApplicationTracking.candidate_id.in_(matching_ids),
            models.ApplicationTracking.job_id == clean_job_id
        ).all()
        tracker_map = {t.candidate_id: t for t in trackers}

        job_posting = db.query(models.JobPosting).filter(models.JobPosting.id == clean_job_id).first()
        min_exp = job_posting.min_experience if job_posting else 1.0

        from utils.scoring import calculate_candidate_score

        result = []
        for c in candidates:
            score_res = calculate_candidate_score(c, min_exp)
            profile_score = score_res["total_score"]
            profile_score_breakdown = score_res["breakdown"]

            track = tracker_map.get(c.id)
            grad = [e for e in c.higher_education if e.level == 'undergrad']
            postgrad = [e for e in c.higher_education if e.level == 'postgrad']
            phd = [e for e in c.higher_education if e.level == 'phd']

            papers_ct = len([p for p in c.publications if p.pub_type == 'paper'])
            books_ct = len([p for p in c.publications if p.pub_type == 'book'])
            chapters_ct = len([p for p in c.publications if p.pub_type == 'chapter'])

            # Find dynamic highest education level
            highest_edu = 'Bachelors'
            if phd: highest_edu = 'PhD'
            elif postgrad: highest_edu = 'Masters'

            d = {
                "id": c.id,
                "full_name": c.full_name,
                "email": c.email,
                "gender": c.gender,
                "state": c.state,
                "age": c.age or (
                    (datetime.date.today().year - c.dob.year - ((datetime.date.today().month, datetime.date.today().day) < (c.dob.month, c.dob.day)))
                    if c.dob else None
                ),
                "years_of_experience": c.years_of_experience,
                "highest_education": highest_edu,
                "current_status": track.current_status if track else 'received',
                "profile_score": profile_score,
                "profile_score_breakdown": profile_score_breakdown,
                "graduation": [{"degree_name": g.degree_name, "university": g.university, "score": f"{g.score_value} {g.score_type}"} for g in grad],
                "postgraduate": [{"degree_name": p.degree_name, "university": p.university, "score": f"{p.score_value} {p.score_type}"} for p in postgrad],
                "doctorate": [{"university": d.university, "thesis_title": d.degree_name, "phd_domain": getattr(d, "phd_domain", None), "grad_year": d.grad_year, "is_pursuing": getattr(d, "is_pursuing", False)} for d in phd],
                "work_experiences": [{"role": w.role, "company_name": w.company_name} for w in c.work_experiences],
                "books_count": books_ct,
                "papers_count": papers_ct,
                "chapters_count": chapters_ct
            }
            result.append(d)
        
        result.sort(key=lambda x: x['full_name'])
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/jobs/{job_id}/filter-options", dependencies=[Depends(get_current_admin)])
def get_filter_options(job_id: str, db: Session = Depends(get_db)):
    clean_id = str(job_id).strip()
    
    # Discover unique values joining application tracking
    states = db.query(models.CandidateMetadata.state).join(
        models.ApplicationTracking, models.ApplicationTracking.candidate_id == models.CandidateMetadata.id
    ).filter(models.ApplicationTracking.job_id == clean_id).distinct().all()

    genders = db.query(models.CandidateMetadata.gender).join(
        models.ApplicationTracking, models.ApplicationTracking.candidate_id == models.CandidateMetadata.id
    ).filter(models.ApplicationTracking.job_id == clean_id).distinct().all()
    
    # Level-specific university discovery
    ug_unis = db.query(models.CandidateHigherEducation.university).join(
        models.ApplicationTracking, models.ApplicationTracking.candidate_id == models.CandidateHigherEducation.candidate_id
    ).filter(
        models.ApplicationTracking.job_id == clean_id,
        models.CandidateHigherEducation.level == 'undergrad'
    ).distinct().all()

    pg_unis = db.query(models.CandidateHigherEducation.university).join(
        models.ApplicationTracking, models.ApplicationTracking.candidate_id == models.CandidateHigherEducation.candidate_id
    ).filter(
        models.ApplicationTracking.job_id == clean_id,
        models.CandidateHigherEducation.level == 'postgrad'
    ).distinct().all()

    phd_unis = db.query(models.CandidateHigherEducation.university).join(
        models.ApplicationTracking, models.ApplicationTracking.candidate_id == models.CandidateHigherEducation.candidate_id
    ).filter(
        models.ApplicationTracking.job_id == clean_id,
        models.CandidateHigherEducation.level == 'phd'
    ).distinct().all()

    return {
        "states": sorted([s[0] for s in states if s[0]]),
        "genders": sorted([g[0] for g in genders if g[0]]),
        "ug_unis": sorted([u[0] for u in ug_unis if u[0]]),
        "pg_unis": sorted([u[0] for u in pg_unis if u[0]]),
        "phd_unis": sorted([u[0] for u in phd_unis if u[0]])
    }

@app.get("/api/v1/jobs/{job_id}/candidates", dependencies=[Depends(get_current_admin)])
def get_job_candidates(job_id: str, db: Session = Depends(get_db)):
    trackers = db.query(models.ApplicationTracking).filter(
        models.ApplicationTracking.job_id == job_id
    ).order_by(models.ApplicationTracking.submitted_at.desc()).all()
    
    if not trackers:
        return []

    job_posting = db.query(models.JobPosting).filter(models.JobPosting.id == job_id).first()
    min_exp = job_posting.min_experience if job_posting else 1.0

    candidate_ids = [t.candidate_id for t in trackers]
    candidates = db.query(models.CandidateMetadata).filter(
        models.CandidateMetadata.id.in_(candidate_ids)
    ).options(
        joinedload(models.CandidateMetadata.higher_education),
        joinedload(models.CandidateMetadata.schooling),
        joinedload(models.CandidateMetadata.publications),
        joinedload(models.CandidateMetadata.work_experiences)
    ).all()
    
    # Map for sorting order preservation
    cand_map = {c.id: c for c in candidates}
    
    from utils.scoring import calculate_candidate_score

    result = []
    for t in trackers:
        c = cand_map.get(t.candidate_id)
        if not c:
            continue
            
        score_res = calculate_candidate_score(c, min_exp)

        grad = [e for e in c.higher_education if e.level == 'undergrad']
        postgrad = [e for e in c.higher_education if e.level == 'postgrad']
        phd = [e for e in c.higher_education if e.level == 'phd']

        papers_ct = len([p for p in c.publications if p.pub_type == 'paper'])
        books_ct = len([p for p in c.publications if p.pub_type == 'book'])
        chapters_ct = len([p for p in c.publications if p.pub_type == 'chapter'])

        highest_edu = 'Bachelors'
        if phd: highest_edu = 'PhD'
        elif postgrad: highest_edu = 'Masters'

        d = {
            "id": c.id,
            "full_name": c.full_name,
            "email": c.email,
            "gender": c.gender,
            "state": c.state,
            "age": c.age or (
                (datetime.date.today().year - c.dob.year - ((datetime.date.today().month, datetime.date.today().day) < (c.dob.month, c.dob.day)))
                if c.dob else None
            ),
            "years_of_experience": c.years_of_experience,
            "highest_education": highest_edu,
            "current_status": t.current_status,
            "profile_score": score_res["total_score"],
            "profile_score_breakdown": score_res["breakdown"],
            "ai_match_score": None,
            "graduation": [{"degree_name": g.degree_name, "university": g.university, "score": f"{g.score_value} {g.score_type}"} for g in grad],
            "postgraduate": [{"degree_name": p.degree_name, "university": p.university, "score": f"{p.score_value} {p.score_type}"} for p in postgrad],
            "doctorate": [{"university": d.university, "thesis_title": d.degree_name, "phd_domain": getattr(d, "phd_domain", None), "grad_year": d.grad_year, "is_pursuing": getattr(d, "is_pursuing", False)} for d in phd],
            "work_experiences": [{"role": w.role, "company_name": w.company_name} for w in c.work_experiences],
            "books_count": books_ct,
            "papers_count": papers_ct,
            "chapters_count": chapters_ct
        }
        result.append(d)
    return result


# ─────────────────────────────────────────────
# List Jobs (with application counts)
# ─────────────────────────────────────────────
@app.get("/api/v1/jobs", dependencies=[Depends(get_current_admin)])
def list_jobs(db: Session = Depends(get_db)):
    today = datetime.date.today()

    # Auto-close open jobs whose deadline has passed
    expired_jobs = db.query(models.JobPosting).filter(
        models.JobPosting.status == 'open',
        models.JobPosting.deadline != None,
        models.JobPosting.deadline < today,
        models.JobPosting.is_deleted == False
    ).all()
    if expired_jobs:
        for ej in expired_jobs:
            ej.status = 'closed'
        db.commit()

    jobs = db.query(models.JobPosting).filter(
        models.JobPosting.is_deleted == False
    ).order_by(models.JobPosting.created_at.desc()).all()

    result = []
    for job in jobs:
        app_count = db.query(models.ApplicationTracking).filter(
            models.ApplicationTracking.job_id == job.id
        ).count()

        result.append({
            "id":             job.id,
            "title":          job.title,
            "position":       job.position,
            "division":       job.division,
            "description":    job.description,
            "key_terms":      job.key_terms,
            "requirements":   job.requirements,
            "keywords":       job.keywords,
            "status":         job.status,
            "total_openings": job.total_openings,
            "deadline":       job.deadline.isoformat() if job.deadline else None,
            "created_by":     job.created_by,
            "created_at":     job.created_at.isoformat(),
            "updated_at":     job.updated_at.isoformat(),
            "is_deleted":     job.is_deleted,
            "application_count": app_count,
            "min_pay":        job.min_pay,
            "max_pay":        job.max_pay,
            "min_experience": job.min_experience,
            "max_experience": job.max_experience,
            "contract_period": job.contract_period,
            "job_mode":       job.job_mode,
            "pay_band":       job.pay_band,
            "pay_level":      job.pay_level,
        })
    return result


# ─────────────────────────────────────────────
# Create Job
# ─────────────────────────────────────────────
@app.post("/api/v1/jobs", dependencies=[Depends(get_current_admin)])
def create_job(payload: schemas.JobPostingCreate, db: Session = Depends(get_db)):
    job_data = payload.model_dump(exclude_unset=True)
    new_job = models.JobPosting(**job_data)
    try:
        db.add(new_job)
        db.commit()
        db.refresh(new_job)
        return {"status": "created", "id": new_job.id}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))


# ─────────────────────────────────────────────
# Update Job (Edit)
# ─────────────────────────────────────────────
@app.patch("/api/v1/jobs/{job_id}", dependencies=[Depends(get_current_admin)])
def update_job(job_id: str, payload: schemas.JobPostingCreate, db: Session = Depends(get_db)):
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.is_deleted == False
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    for key, value in payload.model_dump(exclude_unset=True).items():
        setattr(job, key, value)

    # Auto transition status based on deadline when updated
    today = datetime.date.today()
    if job.deadline:
        if job.deadline >= today and job.status == 'closed':
            job.status = 'open'
        elif job.deadline < today and job.status == 'open':
            job.status = 'closed'

    job.updated_at = datetime.datetime.utcnow()
    try:
        db.commit()
        db.refresh(job)
        return {"status": "updated", "id": job.id}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))



# ─────────────────────────────────────────────
# Publish Draft → Open
# ─────────────────────────────────────────────
@app.patch("/api/v1/jobs/{job_id}/publish", dependencies=[Depends(get_current_admin)])
def publish_job(job_id: str, db: Session = Depends(get_db)):
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.is_deleted == False
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    job.status = 'open'
    job.updated_at = datetime.datetime.utcnow()
    db.commit()
    return {"status": "published"}


# ─────────────────────────────────────────────
# Archive Job
# ─────────────────────────────────────────────
@app.patch("/api/v1/jobs/{job_id}/archive", dependencies=[Depends(get_current_admin)])
def archive_job(job_id: str, db: Session = Depends(get_db)):
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.is_deleted == False
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    job.status = 'archived'
    job.updated_at = datetime.datetime.utcnow()
    db.commit()
    return {"status": "archived"}


# ─────────────────────────────────────────────
# Close Job
# ─────────────────────────────────────────────
@app.patch("/api/v1/jobs/{job_id}/close", dependencies=[Depends(get_current_admin)])
def close_job(job_id: str, db: Session = Depends(get_db)):
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.is_deleted == False
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    job.status = 'closed'
    job.updated_at = datetime.datetime.utcnow()
    db.commit()
    return {"status": "closed"}


# ─────────────────────────────────────────────
# Soft Delete Draft
# ─────────────────────────────────────────────
@app.delete("/api/v1/jobs/{job_id}", dependencies=[Depends(get_current_admin)])
def delete_job(job_id: str, db: Session = Depends(get_db)):
    job = db.query(models.JobPosting).filter(
        models.JobPosting.id == job_id,
        models.JobPosting.status == 'draft',
        models.JobPosting.is_deleted == False
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Draft not found or not deletable")
    job.is_deleted = True
    job.updated_at = datetime.datetime.utcnow()
    db.commit()
    return {"status": "deleted"}





